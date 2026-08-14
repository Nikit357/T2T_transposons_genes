#!/usr/bin/env python
"""Build Extensive_discussion_260803.docx from the passages moved out of the Discussion.

Reviewer major comment 6 asks for a substantially shorter Discussion. Decision D7 keeps the
excised material rather than discarding it, as a supplementary file. The build procedure is
the one plan section 3.5 requires and WP8 spells out: **copy the manuscript and delete**,
never create a fresh document and paste. The copy carries
`word/webextensions/webextension1.xml` with it, and with it the MENDELEY_CITATIONS payload,
so the citations inside the moved paragraphs stay live and Mendeley can refresh a
bibliography containing only the references this file still cites. A pasted reconstruction
would produce dead plain-text citation numbers.

Deletions here are hard, not tracked: this is a new standalone document, not a revision of
one the journal has seen. The tracked deletions of the same passages live in the main
manuscript, applied by 13_manuscript_tracked_edits.py.

    Revised_manuscript/T2T_genes_article_G3_submitted_baseline_260418.docx  (input, never
                                                                            modified)
        -> Revised_manuscript/Extensive_discussion_260803.docx

The baseline rather than the revision is the source, because in the revision these same
paragraphs are already marked deleted.

Usage
    ~/venvs/collagen_3_11/bin/python 14_build_extensive_discussion.py
"""

import os
import re
import shutil
import sys

sys.path.insert(0, os.path.expanduser("~/.claude/skills"))

import docx  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from word_rewrite_trackchanges import (  # noqa: E402
    find_all_p, heading, ins_paragraph, is_p, make_run, set_revision_identity, style_of,
    text_of,
)

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
# All three manuscript docx files live together in revision_G3/Revised_manuscript/.
DOCX_DIR = os.path.join(HERE, "Revised_manuscript")
BASELINE = os.path.join(DOCX_DIR, "T2T_genes_article_G3_submitted_baseline_260418.docx")
TARGET = os.path.join(DOCX_DIR, "Extensive_discussion_260803.docx")

SOFT_MAP = {"\xa0": " ", "‑": "-", "–": "-", "—": "-", "−": "-",
            "‘": "'", "’": "'", "“": '"', "”": '"'}
SOFT_DROP = {"​", "‌", "‍", "﻿"}


def soft(text):
    return "".join(SOFT_MAP.get(c, c) for c in text if c not in SOFT_DROP)


# The paragraphs that move here, in the order they will appear. Each entry is a distinctive
# opening substring of a paragraph in the baseline.
MOVED = {
    "window": [
        "In the present study we used a co-mapping window of 10 kb, 5 kb upstream",
        "The current analysis validates this approach by highlighting the importance",
    ],
    "mechanisms": [
        "As a necessary preliminary analysis, we studied degree of enrichment of TE classes",
        "The enrichment of SINEs and SVAs near TSS (1.468- and 1.368-fold enrichment",
        "Our results of SVA elements being enriched near TSS are connected also to the fact",
        "Moreover, SVA elements are the youngest TE class in humans with active evolutionary",
        "DNA transposons (e.g., TcMar, hAT) utilize a \"cut-and-paste\" transposase",
        "LTR elements, derived from ancient endogenous retroviruses, are also heavily depleted",
        "At the level of TE families, our results reveal a more complex regulatory landscape.",
        "The CR1 family (Chicken repeats, LINE) also shows significant enrichment",
        "Whereas SINEs showed 1.468-fold enrichment near TSS, their major families",
        "Finally, we observed no connection of family enrichment near TSS with their "
        "divergence",
    ],
    "functional": [
        "On average, according to our results, each human gene harbors 15.05 TEs",
        "Genes involved in embryogenesis, transcription, and nervous system development",
        "Conversely, genes enriched with TEs (the \"TE top\" group) are involved in RNA "
        "splicing",
        "RNA splicing: Alu elements were early shown to profoundly impact RNA splicing",
        "DNA repair: The enrichment of TEs (particularly Alu elements) near DNA repair genes",
        "Olfactory receptors: L1 elements are specifically enriched near olfactory receptor",
        "Other groups: metals metabolism, other (specific) metabolism, cytoskeleton, cell "
        "death",
    ],
    "families": [
        "At the level of TE classes, the previous landmark study by",
        "At the level of families, Alu elements are of particular interest.",
        "Alus enrichment near genes of RNA processing is corroborated by the recent evidence",
        "The MIR (Mammalian-wide Interspersed Repeat) family, though less numerous than Alus",
        "SVA elements were previously shown to impact host genes via premature termination",
        "Among the TE classes, only LTR families all demonstrated functional associations",
        "DNA transposons, which have been extinct in the human lineage for millions of "
        "years, are mostly functionally neutral",
        "It is important to note than the fact that certain TE families, precisely 30 out of "
        "44",
        "At the level of lowest and highest divergence by TE classes, ancient "
        "high-divergence",
    ],
    "ifna": [
        "One of the most biologically significant findings that we can report is",
        "The enrichment of young LINEs in this region suggests an ongoing evolutionary arms "
        "race.",
    ],
    "cancer": [
        "The contemporary anticancer treatment heavily relies on the single molecule based",
    ],
}

TITLE = ("Extended discussion for: Telomere-to-telomere co-mapping of transposable elements "
         "and human genes identifies a cluster of young L1 elements in the interferon-alpha "
         "domain")

PREAMBLE = [
    "Daniil Nikitin",
    "Institute of Molecular Biology, National Academy of Science of the Republic of Armenia",
    "This file extends the Discussion of the main article. In response to peer review the "
    "Discussion was shortened and refocused on four things: what we found, how it compares "
    "with prior work, what the design cannot show, and which hypotheses it generates. That "
    "left a body of material which is not needed to follow the argument of the paper but "
    "which supports it, documents the literature it rests on, and would be tedious to "
    "reconstruct from the citations alone. It is collected here rather than deleted.",
    "Five things are gathered, in the order a reader of the main article would want them. "
    "The first is the case for the 10 kb proximity window, set out against the windows other "
    "groups have used, since the choice governs every number in the paper and is a convention "
    "rather than a measurement. The second is the mechanistic review of why each TE class "
    "sits where it does relative to transcription start sites: insertion machinery, selection, "
    "neutral mechanisms of entrenchment, and the timing of germline activation. The third is "
    "the functional interpretation of the gene groups that are richest and poorest in TEs. The "
    "fourth is the family-level review, which is where the specific literature on Alu, MIR, "
    "SVA, LTR and DNA transposon function lives, together with the interferon-alpha domain in "
    "its wider immunological context. The fifth is the case for turning TE-defined gene sets "
    "into compound expression signatures for cancer cohorts, which is a research programme in "
    "its own right rather than a conclusion of this paper.",
    "Every statement here is the text of the submitted manuscript with its citations intact; "
    "nothing has been added to the science. Section openings and closings have been written "
    "to connect the passages, and are marked as such where they are new. Numbers quoted in "
    "this file follow the submitted version; where the revision changed a number, the main "
    "article is authoritative.",
]

SECTIONS = [
    ("window", "1. Why a 10 kb window, and what other studies have used",
     "The main article states the window and reports the sensitivity analysis at 5 and 20 kb. "
     "The reasoning behind the choice, and the range of windows in the literature it is drawn "
     "from, are set out here.",
     "The practical consequence is the one the main article acts on: because the window is a "
     "convention, results are reported at three window sizes and the ordering of TE groups is "
     "shown to be robust to the choice."),
    ("mechanisms", "2. Why each TE class sits where it does relative to transcription start "
                   "sites",
     "The main article gives four mechanisms as a framework and a schematic. This section is "
     "the evidence behind them, class by class, and is the part of the discussion most "
     "directly grounded in other people's experimental work rather than in our measurements.",
     "Taken together these accounts explain the direction of the enrichment and depletion "
     "values without requiring any single mechanism to dominate, which is why the main "
     "article presents them as a framework rather than as a conclusion."),
    ("functional", "3. The gene groups richest and poorest in transposable elements",
     "The main article reports the segregation of biological processes between TE-rich and "
     "TE-poor genes and shows it in the functional-group figures. The interpretation of that "
     "segregation, process by process, follows.",
     "The pattern that emerges is a division of labour: processes that require transcriptional "
     "precision are protected from TE insertion, and processes that benefit from combinatorial "
     "novelty tolerate or exploit it."),
    ("families", "4. Family-level function, and the interferon-alpha domain in context",
     "This is the longest section and the one with the most specific literature attached. It "
     "reviews what is known about the function of the individual families our analysis "
     "highlights, and places the interferon-alpha result in the wider context of type I "
     "interferon biology and autoimmune disease.",
     "The family-level picture is more specific than the class-level one and correspondingly "
     "more fragile: several of these associations rest on a handful of genes, which is why the "
     "main article reports the number of terms per family and the fraction of families with "
     "no significant terms at all."),
    ("cancer", "5. From TE-defined gene sets to compound signatures for cancer cohorts",
     "The main article retains the connection between TE de-repression and cancer in "
     "condensed form. The argument for building compound expression signatures from "
     "TE-defined gene sets, and the biomarker literature it rests on, is a programme for "
     "future work rather than a result of this paper, and is set out here.",
     "The step this suggests is concrete: take the gene sets defined by enrichment for a "
     "given TE class or family, measure them as signatures in public tumour cohorts, and ask "
     "whether the evolutionary grouping carries prognostic information that conventional "
     "functional grouping does not."),
]


def build():
    shutil.copyfile(BASELINE, TARGET)
    document = docx.Document(TARGET)
    body = document.element.body
    set_revision_identity("Claude (extended discussion)", "2026-08-04T00:00:00Z")

    # Locate every paragraph to keep, before anything is removed, so lxml references stay
    # valid across the deletions (word_rewrite principle 4).
    keep = {}
    missing = []
    for section, needles in MOVED.items():
        keep[section] = []
        for needle in needles:
            target = soft(needle)
            hits = find_all_p(body, lambda text, t=target: t in soft(text))
            if hits:
                keep[section].append(hits[0])
            else:
                missing.append(needle)
    kept = {id(p) for group in keep.values() for p in group}

    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sect_pr:
            continue
        if is_p(child) and id(child) in kept:
            continue
        if child.tag == qn("w:sdt"):
            # The MENDELEY_BIBLIOGRAPHY control is kept so Mendeley can rebuild a
            # bibliography for exactly the references this file still cites.
            tags = [t.get(qn("w:val")) or "" for t in child.findall(".//" + qn("w:tag"))]
            if any("MENDELEY_BIBLIOGRAPHY" in tag.upper() for tag in tags):
                continue
        body.remove(child)

    # Re-emit the kept paragraphs under their new headings, in the planned order.
    first = body[0] if len(body) else None
    front = [ins_paragraph(TITLE, bold=True)]
    front += [ins_paragraph(text) for text in PREAMBLE]
    for section, title, opening, closing in SECTIONS:
        front.append(heading(title, 1))
        front.append(ins_paragraph(opening, italic=True))
        front += keep.get(section, [])
        if section == "families":
            front += keep.get("ifna", [])
        front.append(ins_paragraph(closing, italic=True))
    for element in front:
        if sect_pr is not None:
            sect_pr.addprevious(element)
        else:
            body.append(element)

    flatten_insertions(body)
    fix_narrative_citations(body)
    document.save(TARGET)
    return missing


# G1b. Two sentences read ungrammatically because a numbered citation is used as the subject's
# name: "A landmark study by (10) that showed…". Both moved here from the main Discussion in
# Phase 5, so this is where they have to be fixed.
#
# The plan's instruction was to substitute the author name ("by Lu et al. (2020)"), but that
# only works after the Mendeley style switch and would then read "by Lu et al. (Lu et al.
# 2020)" once the control renders author-year. Recasting the sentence so the citation is
# parenthetical instead of narrative is grammatical under *both* styles and therefore does not
# depend on Daniil's switch at all.
NARRATIVE_CITATIONS = [
    ("A landmark study by ", "A landmark study "),
    # The citation sits mid-sentence, so removing "by" alone leaves "A landmark study (10)
    # that showed X, utilized Y" — a relative clause followed by a main verb. The clause has
    # to lose its "that" and gain a conjunction for the sentence to close properly.
    (" that showed enrichment of LINEs and LTRs near duplicated genes and SINEs and DNA "
     "elements near singleton genes, utilized",
     " showed enrichment of LINEs and LTRs near duplicated genes and SINEs and DNA elements "
     "near singleton genes, and utilized"),
    ("At the level of TE classes, the previous landmark study by ",
     "At the level of TE classes, a previous landmark study "),
]


def fix_narrative_citations(body):
    """Recast narrative numbered citations so they read correctly in either style."""
    applied = 0
    for old, new in NARRATIVE_CITATIONS:
        target = soft(old)
        for paragraph in body.iter(qn("w:p")):
            for run in paragraph.findall(qn("w:r")):
                node = run.find(qn("w:t"))
                if node is None or not node.text:
                    continue
                if target in soft(node.text):
                    position = soft(node.text).index(target)
                    node.text = (node.text[:position] + new
                                 + node.text[position + len(old):])
                    applied += 1
                    break
            else:
                continue
            break
    print(f"  G1b: recast {applied} narrative citation(s)")
    return applied


def flatten_insertions(body):
    """Turn the tracked insertions used to build this file into ordinary text.

    The helpers emit `<w:ins>` wrappers because they exist for revising a document. This is
    a new standalone file, so its own text must be plain: nothing here is a revision of
    anything the journal has seen, and leaving revision marks would show the whole file as
    an unaccepted insertion.
    """
    for marker in body.findall(".//" + qn("w:ins")):
        parent = marker.getparent()
        if parent.tag == qn("w:rPr"):
            parent.remove(marker)
            continue
        index = list(parent).index(marker)
        for offset, child in enumerate(list(marker)):
            parent.insert(index + offset, child)
        parent.remove(marker)


missing = build()
document = docx.Document(TARGET)
words = sum(len(re.findall(r"\S+", p.text)) for p in document.paragraphs)
print(f"wrote {os.path.relpath(TARGET, REPO)}")
print(f"  {len(document.paragraphs)} paragraphs, {words} words")
for needle in missing:
    print(f"  NOT FOUND: {needle[:74]}")
if missing:
    sys.exit(1)
