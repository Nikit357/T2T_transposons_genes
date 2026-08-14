#!/usr/bin/env python
"""Apply every Phase 5 manuscript edit to the G3 revision docx as tracked changes.

Reads the untouched submitted baseline and writes the revision file, so the script is
idempotent: re-running it never double-applies an edit, and the tracked diff is always
against exactly what the journal received.

    Revised_manuscript/T2T_genes_article_G3_submitted_baseline_260418.docx  (input, never
                                                                            modified)
        -> Revised_manuscript/T2T_genes_article_G3_revision_260803.docx     (output,
                                                                            overwritten)

Accept All in Word yields the clean revision; Reject All restores the baseline.

Stages, in the order of plan section 7.1a
    D  Methods: N = 500 with the length-bias justification, GO FDR 0.05, the window
       rationale moved down from the Discussion, the 5/20 kb and 10 % sensitivity methods,
       the "FDR-adjusted unless stated" sentence, and the canonical input files
    E  Table 1 -> Tables 1 and 2, and the malformed scientific notation
    F  Results numbers that move under FDR 0.05
    G  Results: three new paragraphs (interferon-alpha domain, sensitivity, Lu et al.)
    H  Title and Abstract
    I  Introduction
    J  Discussion restructured into six subsections
    K  Data availability

Why the two custom helpers below exist
    The skill's `tracked_replace` rebuilds a paragraph from its concatenated text, which
    would destroy any `<w:sdt>` Mendeley citation control inside it -- the exact failure
    mode plan section 3.5 warns about. `tracked_replace_safe` instead edits only contiguous
    runs of `<w:r>` siblings and never reaches inside an `<w:sdt>`. `delete_paragraph_safe`
    does the opposite where it is needed: it also marks the runs *inside* citation controls
    as deleted, so accepting the deletion removes the citation with its sentence and
    Mendeley's refresh then drops it from the bibliography.

Usage
    ~/venvs/collagen_3_11/bin/python 13_manuscript_tracked_edits.py [--dry-run]
"""

import copy
import json
import os
import re
import shutil
import sys

sys.path.insert(0, os.path.expanduser("~/.claude/skills"))

import docx  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from word_rewrite_trackchanges import (  # noqa: E402
    delete_paragraph, find_all_p, find_p, heading, ins_paragraph, insert_after, is_p,
    insert_before, make_del_run, make_run, note_paragraph, set_revision_identity, style_of,
    ins_paragraph_runs, text_of, wrap_del, wrap_ins, _rev_attrs,
)

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
# All three manuscript docx files live together in revision_G3/Revised_manuscript/, apart from
# the subfamilies paper they were originally filed with.
DOCX_DIR = os.path.join(HERE, "Revised_manuscript")
BASELINE = os.path.join(DOCX_DIR, "T2T_genes_article_G3_submitted_baseline_260418.docx")
# Daniil edited the revision by hand on 2026-08-04 and accepted part of the tracked diff
# (594 of the 1,124 deletions), so the file below -- not the baseline -- is the current
# state of the paper. This script used to rebuild the target from the baseline with a
# shutil.copyfile, and that copy WAS its idempotency mechanism; running it unchanged
# against the working file would silently discard every acceptance and manual edit. So
# IN_PLACE is the default now, and the copyfile path is reachable only with
# --from-baseline, which also refuses to overwrite the working file.
WORKING = os.path.join(DOCX_DIR, "T2T_genes_article_G3_revision_260804_manual.docx")
SUPERSEDED = os.path.join(DOCX_DIR, "T2T_genes_article_G3_revision_260803.docx")

DRY_RUN = "--dry-run" in sys.argv
FROM_BASELINE = "--from-baseline" in sys.argv
IN_PLACE = not FROM_BASELINE
TARGET = SUPERSEDED if FROM_BASELINE else WORKING

# Every edit carries this identity so Word groups the revision under one author and date.
REVISION_AUTHOR = "Claude (G3 revision)"
REVISION_DATE = "2026-08-04T00:00:00Z"

report = []


def record(stage, label, ok, outcome=None):
    """Log one edit's three-way outcome.

    In place, an edit whose result is ALREADY in the document is the normal case on a
    re-run, not a failure — but an edit whose target text is nowhere to be found still
    has to be a hard error, or the "a silent skip is impossible" guarantee is lost. So
    the outcome is one of `applied` / `already present` / `not found`, and only the last
    makes the run fail.
    """
    if outcome is None:
        outcome = "applied" if ok else "not found"
    report.append({"stage": stage, "edit": label, "applied": bool(ok),
                   "outcome": outcome})
    return ok


def structural_edit_needed(body, stage, label, marker):
    """False when a structural insert has already been made, so it is not made twice.

    `edit()` is naturally idempotent — its search text is gone once it has been replaced.
    Insertions are not: running them again in place would add a second copy of the same
    16 Discussion paragraphs or a third Table 1. Each insertion therefore names a
    `marker` — a distinctive sentence it introduces — and is skipped when that marker is
    already in the document.
    """
    if IN_PLACE and paragraph_by(body, marker[:80]) is not None:
        record(stage, label, True, "already present")
        return False
    return True


def already_present(body, replacements):
    """True when every replacement's NEW text is already in the document.

    Matched on the first 80 softened characters: enough to identify the sentence, short
    enough not to trip over a citation control sitting inside the replacement.
    """
    for _old, new in replacements:
        if not new or not present_in_any_form(body, new):
            return False
    return True


def present_in_any_form(scope, text):
    """True when `text` is in `scope` in its own, house-styled, or renumbered form.

    `scope` may be the body OR a single paragraph, so this reads the element's own text
    instead of going through `find_all_p`, which iterates only DIRECT children and therefore
    sees nothing at all when it is handed a paragraph — the defect that made every
    already-applied check silently return False.
    """
    if is_p(scope):
        haystack = soft(text_of(scope))
    else:
        haystack = "\n".join(soft(text_of(el)) for el in scope if is_p(el))
    for candidate in dict.fromkeys(
            [text, with_house_style_naming(text), with_supplementary_renumbering(text)]):
        if soft(candidate[:80]) in haystack:
            return True
    return False


def replacement_variants(replacements):
    """Each (old, new) pair plus its house-style-renamed form, when that differs."""
    out = []
    for old, new in replacements:
        renamed = (with_house_style_naming(old), with_house_style_naming(new))
        out.append((old, new))
        if renamed[0] != old:
            out.append(renamed)
    return out


# =======================================================================================
# Citation-safe tracked editing
# =======================================================================================
# Word stores typographic variants that a plain-text view of the document hides: non-breaking
# spaces between a term and its citation, zero-width spaces left by the reference manager,
# curly quotes and en dashes. Matching on the literal characters therefore fails for perfectly
# correct target strings, so every comparison below runs on a softened copy and the edit is
# applied to the original characters via an index map.
SOFT_MAP = {"\xa0": " ", "‑": "-", "–": "-", "—": "-", "−": "-",
            "‘": "'", "’": "'", "“": '"', "”": '"', "′": "'"}
SOFT_DROP = {"​", "‌", "‍", "﻿"}


def soften(text):
    """(normalised text, index map back into `text`) for tolerant substring matching."""
    characters, indices = [], []
    for position, character in enumerate(text):
        if character in SOFT_DROP:
            continue
        characters.append(SOFT_MAP.get(character, character))
        indices.append(position)
    return "".join(characters), indices


def soft(text):
    return soften(text)[0]


# 15_house_style.py's G2 sweep renames "Supplementary File 4" to "File S4". Running 13 in
# place against a document 15 has already touched therefore means some of 13's own search
# strings no longer match the text, even though the edit itself is still wanted. This is
# NOT applied inside soften(): that function returns an index map used to splice runs, and
# a rewrite that changes the string length would corrupt it. It is used only to build an
# alternative form of a search string.
SUPPLEMENTARY_FILE = re.compile(r"Supplementary File (\d+)")


def with_house_style_naming(text):
    """The same text as `15_house_style.py` would have left it (G2 file renaming only)."""
    return SUPPLEMENTARY_FILE.sub(r"File S\1", text)


# The supplementary regrouping (decision D-c/D-d): fourteen candidate files became five
# thematic workbooks, so every old label maps to a workbook AND a sheet. One source of truth,
# used by Stage M to renumber the citations and by already_present() to recognise text that has
# already been renumbered. Every old label happens to map to a single target, which is what
# makes a static map correct here: File S1 is cited three times and File S8 twice, and each
# time for the same table.
FILE_S_LABEL = re.compile(r"File S[1-8]\b")
FILE_S_LABEL_MAP = {
    "File S1": "File S1, sheet TSS_TE_intersections",
    "File S2": "File S1, sheet enrichment_families",
    "File S3": "File S2, sheet by_TE_group",
    "File S4": "File S3, sheet GO_TE_groups",
    "File S5": "File S2, sheet by_divergence",
    "File S6": "File S3, sheet GO_by_divergence",
    "File S7": "File S2, sheet by_family",
    "File S8": "File S3, sheet GO_by_family",
}


def with_supplementary_renumbering(text):
    """The same text as Stage M would have left it (house-style rename, then renumbering).

    Needed so that an edit written against the baseline can still tell that its result is
    already in the document after Stage M has renumbered the file it cites.

    One regex pass with a callback, NOT a loop of str.replace: the mapping is self-referential
    (`File S4` -> `File S3, sheet GO_TE_groups`), so a second rule would rewrite the first
    rule's output and turn that into `File S2, sheet by_TE_group, sheet GO_TE_groups`.
    """
    text = with_house_style_naming(text)
    return FILE_S_LABEL.sub(lambda m: FILE_S_LABEL_MAP.get(m.group(0), m.group(0)), text)


def strip_proof_errors(paragraph):
    """Drop <w:proofErr> markers, which otherwise fragment otherwise-contiguous runs.

    They are spell- and grammar-check hints with no content; Word regenerates them.
    """
    for marker in paragraph.findall(qn("w:proofErr")):
        paragraph.remove(marker)


def run_blocks(paragraph):
    """Maximal contiguous groups of <w:r> children, split by anything else.

    A `<w:sdt>` citation control therefore always terminates a block, which is what keeps
    the rewrite below from ever touching one.
    """
    blocks, current = [], []
    for child in list(paragraph):
        if child.tag == qn("w:r"):
            current.append(child)
        elif current:
            blocks.append(current)
            current = []
    if current:
        blocks.append(current)
    return blocks


def block_text(block):
    return "".join(t.text or "" for r in block for t in r.findall(qn("w:t")))


def tracked_replace_safe(paragraph, replacements):
    """Apply (old, new) substring edits as tracked del+ins without touching citations.

    Each replacement is applied inside the first run-block that contains the whole target
    string; `new=""` is a pure deletion. Returns the list of targets that were not found so
    the caller can fail loudly rather than silently skip an edit.
    """
    strip_proof_errors(paragraph)
    missed = []
    for old, new in replacements:
        placed = False
        target = soft(old)
        for block in run_blocks(paragraph):
            text = block_text(block)
            normalised, index_map = soften(text)
            position = normalised.find(target)
            if position < 0:
                continue
            start = index_map[position]
            end = index_map[position + len(target) - 1] + 1
            before, matched, after = text[:start], text[start:end], text[end:]
            base = block[0].find(qn("w:rPr"))
            pieces = []
            if before:
                pieces.append(make_run(before, base=base))
            pieces.append(wrap_del(make_del_run(matched, base=base)))
            if new:
                pieces.append(wrap_ins(make_run(new, base=base)))
            if after:
                pieces.append(make_run(after, base=base))
            anchor = block[-1]
            for piece in pieces:
                anchor.addprevious(piece)
            for run in block:
                paragraph.remove(run)
            placed = True
            break
        if not placed:
            missed.append(old)
    return missed


def delete_paragraph_safe(paragraph):
    """Tracked-delete a paragraph, including the runs inside its citation controls."""
    delete_paragraph(paragraph)
    for sdt in paragraph.findall(".//" + qn("w:sdt")):
        for run in sdt.findall(".//" + qn("w:r")):
            if run.getparent().tag == qn("w:del"):
                continue
            for text in run.findall(qn("w:t")):
                text.tag = qn("w:delText")
            marker = _rev_attrs(OxmlElement("w:del"))
            run.getparent().replace(run, marker)
            marker.append(run)


def replace_inside_insertions(scope, pairs):
    """Apply (old, new) to text inside `<w:ins>` runs within `scope`, editing in place.

    Needed because `run_blocks` deliberately cannot see inside an insertion — which is what
    protects the Mendeley citation controls, but also hides every paragraph Phase 5 added.
    Editing our own inserted text directly is the correct tracked semantics: the span is
    already marked as added, so there is nothing to track a second time, and Reject All still
    restores the baseline because the whole run disappears with the insertion.

    Returns the per-pair hit counts. Longest target first, so a short target cannot eat a
    prefix of a longer one.
    """
    counts = {old: 0 for old, _ in pairs}
    for old, new in sorted(pairs, key=lambda pair: -len(pair[0])):
        target = soft(old)
        for insertion in list(scope.iter(qn("w:ins"))):
            for run in insertion.findall(qn("w:r")):
                node = run.find(qn("w:t"))
                if node is None or not node.text:
                    continue
                # One left-to-right pass per run, advancing past each replacement. NOT a
                # restart-from-the-beginning loop: `new` legitimately contains `old` when an
                # edit appends to a sentence (Figure 7's caption keeps its existing text and
                # adds the panel statistics after it), and re-scanning from the start would
                # then match the text just written and duplicate it until the loop bound.
                text, offset = node.text, 0
                while True:
                    normalised, index_map = soften(text[offset:])
                    position = normalised.find(target)
                    if position < 0:
                        break
                    start = offset + index_map[position]
                    end = offset + index_map[position + len(target) - 1] + 1
                    text = text[:start] + new + text[end:]
                    offset = start + len(new)
                    counts[old] += 1
                node.text = text
    return counts


def edit_insertion(body, stage, anchor, replacements, occurrence=0):
    """Stage-M edit whose target text lives inside a `<w:ins>` (Phase 5's own text)."""
    paragraph = paragraph_by(body, anchor, occurrence)
    if paragraph is None:
        if IN_PLACE and already_present(body, replacements):
            for old, _new in replacements:
                record(stage, f"ins-replace: {old[:60]}", True, "already present")
            return True
        return record(stage, f"locate for ins-replace: {anchor[:60]}", False)
    # An append-style edit (new begins with old, e.g. Figure 7's caption keeps its sentence
    # and adds the panel statistics after it) leaves its own search string in place, so
    # without this pre-check a second run would append a second copy.
    pending, skipped = [], []
    for old, new in replacements:
        if new and new != old and present_in_any_form(paragraph, new):
            skipped.append(old)
        else:
            pending.append((old, new))
    for old in skipped:
        record(stage, f"ins-replace: {old[:60]}", True, "already present")
    if not pending:
        return True
    replacements = pending
    counts = replace_inside_insertions(paragraph, replacements)
    unresolved = False
    for old, new in replacements:
        if counts[old]:
            record(stage, f"ins-replace: {old[:60]}", True, "applied")
        elif new and paragraph_by(body, new[:80]) is not None:
            record(stage, f"ins-replace: {old[:60]}", True, "already present")
        else:
            record(stage, f"ins-replace: {old[:60]}", False)
            unresolved = True
    return not unresolved


def renumber_citations_in(body, stage, anchor, pairs, occurrence=0):
    """Renumber `File Sn` citations inside one paragraph, via a sentinel so no pair can
    rewrite another pair's output.

    Without the sentinel this is genuinely unsafe: `File S3` -> `File S2, sheet by_TE_group`
    produces text containing `File S2`, which is itself an old target, so a second pair would
    rewrite it again into `File S1, sheet enrichment_families, sheet by_TE_group`.
    """
    paragraph = paragraph_by(body, anchor, occurrence)
    if paragraph is None:
        return record(stage, f"locate for renumber: {anchor[:60]}", False)
    # Per-pair skip FIRST. `File S1` -> `File S1, sheet TSS_TE_intersections` is a
    # self-containing mapping: its result still contains its search string, so a second run
    # would produce `File S1, sheet TSS_TE_intersections, sheet TSS_TE_intersections`.
    pending = []
    for old, new in pairs:
        if present_in_any_form(paragraph, new):
            record(stage, f"renumber: {old} @ {anchor[:40]}", True, "already present")
        else:
            pending.append((old, new))
    if not pending:
        return True
    pairs = pending
    # A plain-ASCII sentinel, not a control character: `soften()` normalises the text
    # it searches, and a sentinel it might strip would make stage 2 find nothing.
    sentinels = [(old, f"@@FS{index}@@") for index, (old, _new) in enumerate(pairs)]
    first = replace_inside_insertions(paragraph, sentinels)
    if not any(first.values()):
        # Already renumbered: the new form is present and the old label is gone.
        if all(paragraph_by(body, new[:60]) is not None for _old, new in pairs):
            for old, _new in pairs:
                record(stage, f"renumber: {old} @ {anchor[:40]}", True, "already present")
            return True
        for old, _new in pairs:
            record(stage, f"renumber: {old} @ {anchor[:40]}", False)
        return False
    replace_inside_insertions(
        paragraph, [(sentinel, new) for (old, sentinel), (_o, new)
                    in zip(sentinels, pairs)])
    unresolved = False
    for (old, _sentinel), (_o, new) in zip(sentinels, pairs):
        if first[old]:
            record(stage, f"renumber: {old} -> {new} @ {anchor[:40]}", True, "applied")
        else:
            record(stage, f"renumber: {old} @ {anchor[:40]}", False)
            unresolved = True
    return not unresolved


def caption_like(text, model):
    """A tracked-inserted caption paragraph formatted like the captions around it.

    The existing supplementary captions carry the `Strong` character style on their runs,
    which is what makes the `Figure Sn.` label bold (G2). `ins_paragraph` sets the paragraph
    style only, so a new caption built with it would sit among them unstyled; cloning the
    model caption's run properties matches them exactly instead of guessing at bold.
    """
    base = None
    for run in model.iter(qn("w:r")):
        node = run.find(qn("w:t")) if run.find(qn("w:t")) is not None else run.find(qn("w:delText"))
        if node is not None and (node.text or "").strip():
            base = run.find(qn("w:rPr"))
            break
    return ins_paragraph_runs([make_run(text, base=base)], style=style_of(model))


def delete_paragraph_fully(paragraph):
    """Tracked-delete a paragraph that also contains our own `<w:ins>` runs.

    An insertion that is later removed by the same revision is simply dropped: it was never
    in the baseline, so there is nothing for a reviewer to accept or reject, and leaving it
    behind would show the paragraph as half struck-through and half live. The original runs
    are tracked-deleted as usual, so Reject All still restores the baseline paragraph.
    """
    for insertion in list(paragraph.iter(qn("w:ins"))):
        parent = insertion.getparent()
        if parent is not None:
            parent.remove(insertion)
    delete_paragraph_safe(paragraph)


def paragraph_by(body, needle, occurrence=0):
    """The nth body paragraph whose visible text contains `needle`, softly matched.

    Tried in both the pre- and post-house-style naming of the supplementary files, so a
    needle written against the baseline still locates its paragraph after 15 has run.
    """
    for candidate in dict.fromkeys([needle, with_house_style_naming(needle)]):
        target = soft(candidate)
        hits = find_all_p(body, lambda text: target in soft(text))
        if len(hits) > occurrence:
            return hits[occurrence]
    return None


def edit(body, stage, needle, replacements, occurrence=0):
    paragraph = paragraph_by(body, needle, occurrence)
    if paragraph is None:
        if IN_PLACE and already_present(body, replacements):
            for old, _new in replacements:
                record(stage, f"replace: {old[:70]}", True, "already present")
            return True
        return record(stage, f"locate: {needle[:70]}", False)
    pending, skipped = [], []
    for old, new in replacements:
        if new and new != old and present_in_any_form(paragraph, new):
            skipped.append(old)
        else:
            pending.append((old, new))
    for old in skipped:
        record(stage, f"replace: {old[:70]}", True, "already present")
    if not pending:
        return True
    replacements = pending
    missed = tracked_replace_safe(paragraph, replacements)
    if missed:
        # Retry every miss in the post-house-style naming before calling it a failure.
        retry = [(with_house_style_naming(old), with_house_style_naming(new))
                 for old, new in replacements
                 if old in missed and with_house_style_naming(old) != old]
        if retry:
            still_missed = tracked_replace_safe(paragraph, retry)
            recovered = {old for old, _new in replacements
                         if with_house_style_naming(old) not in still_missed
                         and with_house_style_naming(old) != old}
            missed = [old for old in missed if old not in recovered]
    unresolved = False
    for old, new in replacements:
        if old not in missed:
            record(stage, f"replace: {old[:70]}", True, "applied")
        elif IN_PLACE and new and paragraph_by(body, new[:80]) is not None:
            record(stage, f"replace: {old[:70]}", True, "already present")
        else:
            record(stage, f"replace: {old[:70]}", False)
            unresolved = True
    return not unresolved


def insert_paragraphs_after(body, stage, needle, texts, style=None, occurrence=0):
    paragraph = paragraph_by(body, needle, occurrence)
    if paragraph is None:
        return record(stage, f"locate for insert: {needle[:70]}", False)
    if IN_PLACE and texts and all(paragraph_by(body, t[:80]) is not None for t in texts):
        record(stage, f"insert {len(texts)} paragraph(s) after: {needle[:60]}", True,
               "already present")
        # The EXISTING paragraphs, not an empty list: later edits use the return value as
        # an anchor (the Lu et al. paragraph is placed after the robustness paragraphs),
        # and an empty list would make those edits report a false failure.
        return [paragraph_by(body, t[:80]) for t in texts]
    elements = [ins_paragraph(text, style=style or style_of(paragraph)) for text in texts]
    insert_after(paragraph, elements)
    record(stage, f"insert {len(texts)} paragraph(s) after: {needle[:60]}", True)
    return elements


# =======================================================================================
# Tracked tables
# =======================================================================================
def mark_row_inserted(row):
    trpr = row.find(qn("w:trPr"))
    if trpr is None:
        trpr = OxmlElement("w:trPr")
        row.insert(0, trpr)
    trpr.append(_rev_attrs(OxmlElement("w:ins")))


def mark_row_deleted(row):
    trpr = row.find(qn("w:trPr"))
    if trpr is None:
        trpr = OxmlElement("w:trPr")
        row.insert(0, trpr)
    trpr.append(_rev_attrs(OxmlElement("w:del")))


def build_tracked_table(document, header, rows, template=None):
    """A whole table marked as a tracked insertion: rows, paragraph marks and runs."""
    table = document.add_table(rows=1 + len(rows), cols=len(header))
    if template is not None:
        table.style = template
    for column, value in enumerate(header):
        cell = table.cell(0, column)
        cell.text = ""
        paragraph = cell.paragraphs[0]._p
        paragraph.append(wrap_ins(make_run(value, bold=True)))
        _mark_paragraph_inserted(paragraph)
    for index, values in enumerate(rows, start=1):
        for column, value in enumerate(values):
            cell = table.cell(index, column)
            cell.text = ""
            paragraph = cell.paragraphs[0]._p
            paragraph.append(wrap_ins(make_run(str(value))))
            _mark_paragraph_inserted(paragraph)
    element = table._tbl
    for row in element.findall(qn("w:tr")):
        mark_row_inserted(row)
    element.getparent().remove(element)
    return element


def _mark_paragraph_inserted(paragraph):
    ppr = paragraph.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph.insert(0, ppr)
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        ppr.append(rpr)
    if rpr.find(qn("w:ins")) is None:
        rpr.append(_rev_attrs(OxmlElement("w:ins")))


def delete_tracked_table(table_element):
    """Mark every row, cell paragraph and run of an existing table as deleted."""
    for row in table_element.findall(qn("w:tr")):
        mark_row_deleted(row)
        for paragraph in row.findall(".//" + qn("w:p")):
            delete_paragraph_safe(paragraph)


# =======================================================================================
# Numbers the text quotes, read from the persisted outputs
# =======================================================================================
with open(os.path.join(HERE, "output", "results_numbers.json")) as handle:
    numbers = json.load(handle)
with open(os.path.join(HERE, "output", "length_bias_correlation.json")) as handle:
    length_bias = json.load(handle)

# Stage M's numbers, likewise from the persisted outputs rather than typed by hand.
with open(os.path.join(HERE, "output", "network_qc.json")) as handle:
    network_qc = json.load(handle)
with open(os.path.join(HERE, "output", "sankey_ribbon_filter.json")) as handle:
    sankey_filter = json.load(handle)
with open(os.path.join(HERE, "output", "go_grid_summary.json")) as handle:
    go_grid = json.load(handle)
with open(os.path.join(HERE, "supplementary", "INVENTORY.json")) as handle:
    supplementary = json.load(handle)
figure7_stats = {}
with open(os.path.join(HERE, "output", "figure7_statistics.csv")) as handle:
    import csv as _csv
    for row in _csv.DictReader(handle):
        figure7_stats[row["panel"]] = row


def qc_top_n(panel):
    """The GO terms per group a network panel actually achieved (never guessed)."""
    return int(network_qc[f"{panel}.svg"]["top_n"])


def qc_wrap_sentence(panel):
    """The caption sentence about label wrapping, or '' if that panel was not wrapped.

    Whether a panel needed wrapping is a ladder outcome, not a setting — S9 and S10 come
    out clean at half width on the layout-spacing rung alone — so the caption reads it off
    `network_qc.json` rather than asserting it.
    """
    wrap = network_qc[f"{panel}.svg"].get("label_wrap")
    return f"Term names are wrapped at {wrap} characters. " if wrap else ""


def mwu_sentence(panel, first_label, second_label):
    """The Mann-Whitney sentence for a Figure 7 box panel, from figure7_statistics.csv."""
    row = figure7_stats[panel]
    p_raw = float(row["p_raw"])
    p_text = "p < 0.001" if p_raw < 0.001 else f"p = {p_raw:.3f}"
    first, second = row["n"].split(" vs ")
    return (f"Mann-Whitney U, raw {p_text}; n = {first} {first_label} versus "
            f"{second} {second_label}")

ifna_qc = numbers["ifna"]["qc"]
ifna_tests = {row["test"]: row for row in numbers["ifna"]["tests"]}
assembly = numbers["assembly_bound"]


def load_document():
    """Open the document this run edits, and never overwrite work in place of a copy."""
    if FROM_BASELINE:
        if os.path.abspath(TARGET) == os.path.abspath(WORKING):
            raise SystemExit(
                "--from-baseline would overwrite the working file. Refusing: it carries "
                "Daniil's manual edits and his acceptances of the tracked diff."
            )
        if not DRY_RUN:
            shutil.copyfile(BASELINE, TARGET)
        return _opened(TARGET if not DRY_RUN else BASELINE)
    if not os.path.exists(WORKING):
        raise SystemExit(f"the working file is missing: {WORKING}")
    # In place: read and write the same file. The caller is expected to have it in git or
    # in a backup, because there is no pristine copy to fall back to any more.
    return _opened(WORKING)


def _opened(path):
    print(f"editing {'(dry run) ' if DRY_RUN else ''}{os.path.basename(path)} "
          f"[{'in place' if IN_PLACE else 'rebuilt from the baseline'}]")
    document = docx.Document(path)
    return document, document.element.body


document, body = load_document()
set_revision_identity(REVISION_AUTHOR, REVISION_DATE)


# =======================================================================================
# Stage H -- Title and Abstract  (applied first only because both sit at the top of the
# body; the writing order of section 7.1a is preserved in the prose itself)
# =======================================================================================
TITLE_OLD = ("Evolutionary arms race between transposable elements and human genes: "
             "telomere-to-telomere genome comprehensive analysis identifies young L1 "
             "clusters in the interferon-alpha domain")
TITLE_NEW = ("Telomere-to-telomere co-mapping of transposable elements and human genes "
             "identifies a cluster of young L1 elements in the interferon-alpha domain")
edit(body, "H", "Evolutionary arms race between transposable elements",
     [(TITLE_OLD, TITLE_NEW)])

# The abstract must stay within 250 words (G15), so the new interferon-alpha statistics are
# paid for by compressing the framing sentences rather than added on top of them.
edit(body, "H", "Transposable elements (TEs) have played a central role", [
    ("The availability of the complete telomere-to-telomere (T2T) human genome assembly "
     "enables comprehensive investigation of TE contributions to gene regulation. Using a "
     "10 kb window in the T2T genome, we performed comprehensive mapping of 3,709,429 human "
     "TEs to 28,738 genes with random background correction and assessed the enrichment and "
     "functional associations of six TE classes and 44 families.",
     "The complete telomere-to-telomere (T2T) assembly makes it possible to measure where "
     "TEs sit relative to every annotated gene. Using a 10 kb window around transcription "
     "start sites, we mapped 3,709,429 human TEs to 28,738 genes against a 500-permutation "
     "background that corrects the length bias of the odds ratio, and assessed the "
     "enrichment and functional associations of six TE classes and 44 families."),
    ("We identified a 220 kb interferon-alpha genomic domain enriched with evolutionarily "
     "young L1 elements, suggesting a recent evolutionary arms race influencing innate "
     "immune responses.",
     "We identified a 220 kb interferon-alpha domain enriched with young L1 elements: 77 of "
     "its 175 TEs are L1 copies from 36 subfamilies, at mean divergence 135.7 against 188.2 "
     "genome-wide, a deficit that survives 10,000 random windows matched for L1 count "
     "(empirical p = 0.006) or gene density (p = 0.002). This is consistent with recent L1 "
     "activity; whether these elements affect innate immune gene regulation remains "
     "untested."),
    ("MIR elements were associated with genes involved in zinc, copper, and cadmium "
     "detoxification",
     "MIR elements were associated with zinc homeostasis and voltage-gated potassium "
     "channel genes"),
    ("and LTR elements were potentially associated with potassium ion channel function. "
     "This proximity-based analysis provides a foundational framework for evaluating the "
     "functional impact of transposable elements on human gene regulation and their role in "
     "driving regulatory innovation.",
     "and LTR elements were associated with ion transport. This correlative map provides a "
     "positional baseline against which epigenomic and functional TE studies can be "
     "normalised, and a set of candidate loci for follow-up."),
])

# =======================================================================================
# Stage I -- Introduction (WP7)
# =======================================================================================
# The "(4)" that closes this sentence is a Mendeley content control, so the sentence is
# edited on either side of it and the qualifying sentence is added at the head of the run
# block that follows the citation.
edit(body, "I", "TEs proliferate in a wave pattern", [
    ("they are in a continuous evolutionary arms race against the host defense systems ",
     "they have been described as being in a continuous evolutionary arms race with host "
     "defence systems "),
    (". During this process TEs insert their transcription factor binding sites (TFBS) in "
     "the vicinity of host genes and alter their expression ",
     ". Throughout this paper the term \"arms race\" is used only when reporting that cited "
     "literature: the design applied here is correlative and static, and cannot establish an "
     "arms race for our own observations. During this process TEs insert their transcription "
     "factor binding sites (TFBS) in the vicinity of host genes and alter their expression "),
])
edit(body, "I", "There are dozens of studies of this epigenetic impact", [
    ("The unified approach utilizing the most up to date T2T human genome assembly could "
     "significantly improve the overall understanding of TEs-host genome evolutionary arms "
     "race and its impact on regulatory innovation and human health and disease.",
     "A unified proximity map built on the most up-to-date T2T human genome assembly "
     "supplies the positional baseline these studies need: a common measure of how close "
     "each TE group sits to each class of gene, against which epigenomic and expression "
     "signals can be normalised."),
])
edit(body, "I", "Here we have taken the most widely used proximity window", [
    ("LINE elements in general could impact lipid metabolism and sensory perception of "
     "smell, and LTRs are potentially connected with potassium ion channels",
     "LINE elements are enriched near genes of lipid metabolism and sensory perception of "
     "smell, and LTR elements near genes of ion transport"),
    ("We showed that the 220 kb interferon alpha domain is uniquely enriched with young, "
     "low-divergence L1 elements indicating the recent example of evolutionary arms race "
     "shaping innate immune response.",
     "We show that the 220 kb interferon alpha domain is enriched with young, "
     "low-divergence L1 elements, identifying a candidate locus of recent L1 activity for "
     "functional follow-up."),
])

# =======================================================================================
# Stage D -- Methods
# =======================================================================================
# D1. The canonical inputs, named so the repository files can be matched to the text.
insert_paragraphs_after(
    body, "D", "In this study, we adopted the RepeatMasker classification as provided",
    ["The processed annotation used throughout is provided as the tab-separated file "
     "T2T_repeat_masker_processed_sorted.bed (chromosome, start, end, divergence score, "
     "subfamily, family, class), and the transcription start site neighbourhoods as "
     "T2T_genes.bed. These two files are the sole inputs to every intersection reported "
     "here, including the permutation background and the window sensitivity analysis."])

# D2. The window rationale, moved down from the Discussion so the choice is justified where
#     it is made rather than after the fact.
insert_paragraphs_after(
    body, "D", "The mapping of TE coordinates with their divergence onto gene TSS 10 kb",
    ["The 10 kb window (5 kb upstream and 5 kb downstream of each TSS) was chosen because it "
     "is the most widely used proximity window in the functional TE literature and therefore "
     "makes this map comparable to it. A study of interferon-inducible enhancers spread by "
     "LTR elements used a 10 kb window; an analysis of TE content around duplicated and "
     "singleton genes used 4 kb and 20 kb windows; one of the first surveys of TE-linked "
     "transcription factor binding used 5 kb either side of coding exons; a proximity study "
     "of LTR function used a shifted 6 kb window (5 kb upstream, 1 kb downstream); and our "
     "own transcription-factor and histone-modification studies of retroelements used the "
     "same 10 kb window applied here. Because the choice is a convention rather than a "
     "measurement, all key results are additionally reported at 5 kb and 20 kb (below).",
     "Sensitivity to the window and to the gene-set cut-off. The full enrichment analysis "
     "was repeated for TSS neighbourhoods of 5 kb (2.5 kb either side) and 20 kb (10 kb "
     "either side), built from the same TSS definition, with 500 fresh permutations per "
     "window; the 10 kb background is the one described above. The Gene Ontology analysis "
     "was repeated with the top and bottom 10 % of genes (2,872 genes) alongside the 5 % "
     "sets (1,436 genes) used in the main analysis, at the same FDR threshold. Agreement "
     "between conditions was measured by Spearman and Pearson correlation of the observed "
     "to random odds ratio, by Bland-Altman comparison, by the overlap coefficient and a "
     "hypergeometric test for the gene sets, by Kendall tau with a bootstrap confidence "
     "interval for the gene rankings, and by a label-shuffling permutation test of the "
     "observed correlation, so that “concordant” is a measured rather than a "
     "visual statement."])

# D3. N = 500, and why 500 is enough.
edit(body, "D", "To control TE enrichment near genes against a random background", [
    ("we performed 1000 random permutations of TEs coordinates by the bedtools shuffle "
     "command, selecting random state number (seed) from 1 to 1000",
     "we performed 500 random permutations of TE coordinates with the bedtools shuffle "
     "command, using random seeds 1 to 500"),
    ("depending on whether the observed OR was below or above the median across the random "
     "1000 OR values",
     "depending on whether the observed OR was below or above the median across the 500 "
     "random OR values"),
    ("where N was the permutations number (1000)",
     "where N was the permutations number (500), which places the lower bound at 0.004"),
])
insert_paragraphs_after(
    body, "D", "An enrichment score was calculated as a fold change of observed versus",
    ["The permutation background is what makes these enrichment values interpretable, "
     "because it removes a length bias that the Fisher exact test cannot. The probability "
     "that an element intersects a fixed window grows with the length of the element, so the "
     f"random odds ratio scales almost linearly with mean element length across the 44 "
     f"families (Pearson R = {length_bias['families_pearson_r']:.3f}, "
     f"n = {length_bias['families_n']}, mean lengths "
     f"{length_bias['family_mean_length_min_bp']:,.0f}-"
     f"{length_bias['family_mean_length_max_bp']:,.0f} bp): Alu elements average 316 bp and "
     "a random OR of 1.54, whereas L1 elements average 6,357 bp and a random OR of 2.66. "
     "Reporting the observed odds ratio alone would therefore systematically under-call "
     "short elements and over-call long ones, and every enrichment statement in this work is "
     "consequently made on the ratio of the observed to the random odds ratio rather than on "
     "the observed odds ratio itself.",
     "Five hundred permutations are sufficient for that purpose, and this was verified "
     "rather than assumed. Recomputing the random odds ratio from truncated prefixes of the "
     "permutation set shows that by N = 250 the running mean is already within 0.06 standard "
     "deviations of its N = 500 value for the worst-behaved class and within 0.10 standard "
     "deviations for the worst-behaved family, with the standard deviation itself estimated "
     "to within 4 %; at N = 100 the drift is still 0.18 standard deviations. The "
     "convergence trajectories are shown in Supplementary Figure S14 and the checkpoint "
     "values are provided in the repository."])

# D4. GO threshold, and the raw-versus-corrected statement the journal's statistics policy
#     requires.
edit(body, "D", "FDR-corrected Fisher exact test p-value 0.1 was selected as a threshold", [
    ("FDR-corrected Fisher exact test p-value 0.1 was selected as a threshold value for GO "
     "terms.",
     "FDR-corrected Fisher exact test p-value 0.05 was selected as a threshold value for GO "
     "terms, in the main text and in the supplementary material alike; no separate "
     "“suggestive” band is reported."),
])
insert_paragraphs_after(
    body, "D", "FDR-corrected p-value 0.05 was selected as a threshold value for "
               "Mann-Whitney",
    ["Unless stated otherwise, all reported p-values are Benjamini-Hochberg FDR-adjusted, "
     "and p-values that are raw are labelled as such in the figures, tables and captions. "
     "The Kolmogorov-Smirnov comparisons of divergence and length distributions, the Pearson "
     "correlations of GO term counts against enrichment statistics, and the "
     "interferon-alpha domain permutation tests are single tests per panel and are reported "
     "raw. Both the raw and the adjusted value are provided for every test in the "
     "supplementary tables: enrichment_families_with_random.csv and "
     "enrichment_subfamilies_with_random.csv carry Enrichment_p_value and "
     "Enrichment_p_value_adjusted alongside p_raw_empirical and p_adjusted_empirical_bh, "
     "and the Gene Ontology tables carry both P-value and FDR."])

# D5. The interferon-alpha domain test, described where the other tests are described.
insert_paragraphs_after(
    body, "D", "Distributions were compared by Kolmogorov-Smirnov test",
    ["Interferon-alpha domain permutation test. The 220 kb domain on chromosome 9 was tested "
     "against 10,000 random windows of the same length drawn from a pool of 120,000 "
     "autosomal windows (random seed 42). Three nested null models were used, each "
     "progressively better matched to the domain: windows containing at least one L1 "
     "element; windows containing at least 40 L1 elements, which controls for local L1 "
     "density; and windows containing at least 10 annotated genes and at least one L1, which "
     "controls for gene density. In each case the statistic was the mean divergence of the "
     "L1 elements in the window and the two-sided empirical p-value was the fraction of null "
     "windows at least as extreme as the observation. A fourth test compared the "
     "young-primate-specific (L1HS and L1P*) against the older mammalian (L1M*) composition "
     "of the domain with the rest of the genome by Fisher exact test. Leave-one-out means "
     "were computed to confirm that no single element drives the result."])

# =======================================================================================
# Stage E -- Table 1 -> Tables 1 and 2
# =======================================================================================
edit(body, "E", "Table 1 shows assessment of TE enrichment in genes TSS 10 kb", [
    ("Table 1 shows assessment of TE enrichment in genes TSS 10 kb neighborhoods by class.",
     "Tables 1 and 2 show the assessment of TE enrichment in gene TSS 10 kb neighbourhoods "
     "by class: Table 1 reports the counts and the odds ratios, Table 2 the statistical "
     "support for them."),
])

caption = paragraph_by(body, "Table 1. Enrichment of TE classes in gene TSS neighborhoods")
old_table = None
if caption is not None:
    tracked_replace_safe(caption, [
        ("Table 1. Enrichment of TE classes in gene TSS neighborhoods.",
         "Table 1. Enrichment of TE classes in gene TSS 10 kb neighbourhoods."),
    ])
    cursor = caption.getnext()
    while cursor is not None and cursor.tag != qn("w:tbl"):
        cursor = cursor.getnext()
    old_table = cursor
record("E", "locate Table 1 caption and table", old_table is not None)

TABLE_NOTE_MARKER = ("Tables 1 and 2 replace the original 11-column Table 1 "
                     "(reviewer minor comment 4).")

if old_table is not None and not structural_edit_needed(
        body, "E", "insert Tables 1 and 2 and tracked-delete the original",
        TABLE_NOTE_MARKER):
    old_table = None

if old_table is not None:
    table1_rows = [
        ["LINE", "169,930", "1,005,214", "2.13", "0.877"],
        ["LTR", "51,103", "531,410", "1.11", "0.667"],
        ["SINE", "302,480", "1,706,485", "2.25", "1.468"],
        ["DNA", "57,684", "458,177", "1.51", "0.938"],
        ["SVA", "1,170", "6,274", "2.40", "1.368"],
        ["Helitrons", "173", "1,869", "1.07", "0.661"],
    ]
    table2_rows = [
        ["LINE", "< 10⁻²⁰⁰", "2.43 ± 0.009", "0.004", "0.004"],
        ["LTR", "7.6 × 10⁻¹¹²", "1.67 ± 0.010", "0.004",
         "0.004"],
        ["SINE", "< 10⁻²⁰⁰", "1.53 ± 0.005", "0.004", "0.004"],
        ["DNA", "< 10⁻²⁰⁰", "1.61 ± 0.010", "0.004", "0.004"],
        ["SVA", "1.4 × 10⁻¹³²", "1.75 ± 0.094", "0.004",
         "0.004"],
        ["Helitrons", "0.413", "1.61 ± 0.163", "0.004", "0.004"],
    ]
    template = old_table.find(qn("w:tblPr"))
    new_table_1 = build_tracked_table(
        document, ["Class", "TEs in TSS windows", "TEs total", "Observed OR",
                   "Observed / random OR"], table1_rows)
    caption_2 = ins_paragraph(
        "Table 2. Statistical support for TE class enrichment in gene TSS 10 kb "
        "neighbourhoods. Fisher exact p-values are FDR-adjusted; the empirical p-value is "
        "derived from 500 permutations and its lower bound is 2/501 = 0.004.",
        style=style_of(caption))
    new_table_2 = build_tracked_table(
        document, ["Class", "Adjusted Fisher p", "Random OR (mean ± SD)",
                   "Empirical p", "Adjusted empirical p"], table2_rows)
    note = note_paragraph(
        "Tables 1 and 2 replace the original 11-column Table 1 (reviewer minor comment 4). "
        "The values are unchanged; the unadjusted Fisher p-values moved to "
        "TableS_class_enrichment_full.csv in the supplementary material. Apply the journal's "
        "table style to both when formatting.")
    if template is not None:
        for element in (new_table_1, new_table_2):
            element.insert(0, copy.deepcopy(template))
    insert_after(old_table, [new_table_1, caption_2, new_table_2, note])
    delete_tracked_table(old_table)
    record("E", "insert Tables 1 and 2 and tracked-delete the original", True)

# =======================================================================================
# Stage F -- Results numbers that move under FDR 0.05
# =======================================================================================
# Class level, Figure 4 (line 93 of the plain-text view).
edit(body, "F", "Each gene set was tested against the canonical Gene Ontology", [
    ("(FDR threshold of 0.1 was applied, Supplementary File 4)",
     "(FDR threshold of 0.05 was applied, Supplementary File 4)"),
    # The exponents are superscript-formatted plain digits in the docx, not Unicode
    # superscripts, so the target text reads "10-40 - 10-80".
    ("Moreover, the three all TEs depleted terms with the lowest FDR-corrected p-values in "
     "the entire set (p-values in the range 10-40 – 10-80) related to transcriptional "
     "activators.",
     "Moreover, the three TE-depleted terms with the lowest FDR-corrected p-values in the "
     "entire set (chromatin, 2.3 × 10⁻⁹¹; DNA-binding transcription "
     "factor activity, RNA polymerase II-specific, 3.3 × 10⁻⁸³; and "
     "sequence-specific double-stranded DNA binding, 2.7 × 10⁻⁸¹) all "
     "relate to transcriptional regulators."),
    ("The rest classes with significant GO terms based on genes with highest TE count led to "
     "olfactory receptor activity, flavone metabolism and regulation of fatty acids "
     "metabolism (LINEs), glutaminergic synapse, lipopolysaccharide mediated signaling and "
     "ion transport (LTRs) and RNA polymerase II termination (SVA elements).",
     "The rest of the classes with significant GO terms based on genes with the highest TE "
     "count led to olfactory receptor activity and negative regulation of fatty acid "
     "metabolism (LINEs), ion transport, ion transmembrane transport and synapse (LTRs), "
     "and termination of RNA polymerase II transcription (SVA elements). At the tightened "
     "threshold the LINE class loses flavone metabolic process (FDR = 0.088) and the LTR "
     "class loses glutamatergic synapse and positive regulation of "
     "lipopolysaccharide-mediated signalling (both FDR = 0.086)."),
])
edit(body, "F", "The GO terms were manually (with Gemini pro assistance) classified", [
    ("into 25 major groups", "into 25 major groups, 22 of which are populated at this "
                             "threshold,"),
    ("While there was no systematic co-clustering of molecular processes metagroups, RNA "
     "splicing and processing processes were specifically enriched in SINEs adjacent genes, "
     "metals metabolism was associated with LTRs, DNA replication and recombination related "
     "to all TEs, lipid metabolism was enriched in LINEs adjacent genes.",
     "While there was no systematic co-clustering of molecular process metagroups, RNA "
     "splicing and processing was specifically enriched in SINE-adjacent genes "
     "(FDR = 0.001) and metals metabolism was associated with LTRs (FDR = 0.010). Two "
     "associations that were significant at the previous threshold do not survive at 0.05 "
     "and are no longer claimed: lipid metabolism in LINE-adjacent genes (FDR = 0.021 at "
     "0.1) and DNA replication and recombination in the all-TE top group (FDR = 0.036 at "
     "0.1)."),
])

# Divergence level, Figure 5 (line 103).
edit(body, "F", "We visualized top 30 GO terms of each TE-divergence groups", [
    ("having average divergence of intersecting LINE elements at the level of 95 – "
     "161.7",
     "having per-gene average divergence of the intersecting LINE elements between 95.0 and "
     "161.7, while the mean divergence of the individual L1 elements across the whole domain "
     "is 135.7 against a genome-wide L1 mean of 188.2"),
    ("LTR elements of lowest divergence demonstrated a single term of flavin-based "
     "oxidoreductases.",
     "LTR elements of lowest divergence demonstrated a single term of flavin-dependent "
     "oxidoreductase activity (FDR = 0.039)."),
])

# Family level, Figure 6 (lines 119-137).
edit(body, "F", "GO analysis for each family resulted in only 14 families", [
    ("GO analysis for each family resulted in only 14 families having statistically "
     "significant (FDR-adjusted p-value < 0.1) terms (out of 44 total families): 4 LINE "
     "families (L1, L2, CR1, Dong-R4), 5 LTR ones (ERVL, ERVK, ERV1, ERVL-MaLR and Gypsy), "
     "2 SINE families (Alu and MIR elements), SVA elements and 2 DNA families (hAT-Charlie "
     "and hAT-Tip100) (Figure 6A, Supplementary File 8). 3 of these families (hAT-Tip100, "
     "Dong-R4 and L2) had only 1 significant GO term per family, reflecting likely a random "
     "nature of these enrichments.",
     "GO analysis for each family resulted in only 13 families having statistically "
     "significant (FDR-adjusted p-value < 0.05) terms (out of 44 total families): 3 LINE "
     "families (L1, L2, CR1), 5 LTR ones (ERVL, ERVK, ERV1, ERVL-MaLR and Gypsy), 2 SINE "
     "families (Alu and MIR elements), SVA elements and 2 DNA families (hAT-Charlie and "
     "hAT-Tip100) (Figure 6A, Supplementary File 8). Dong-R4, whose single term the previous "
     "threshold retained, no longer reaches significance, which is consistent with the "
     "random nature we ascribed to it. Two of the remaining families (hAT-Tip100 and L2) "
     "have only 1 significant GO term each, again reflecting the likely random nature of "
     "these enrichments."),
])
edit(body, "F", "Visualization of top 30 GO terms by family in Figure 6A", [
    ("L1 elements, as it was previously shown for LINEs, were connected with olfactory "
     "receptors, fatty acids and flavone metabolism, Dong-R4 was connected with axonal "
     "transport of mitochondrion (3 out of 15 genes), L2 had non-informative protein "
     "binding, CR1 elements had surprising connections with neurotransmitter processes "
     "(postsynaptic receptor internalization and response to cocaine), ventricular septum "
     "development and endopeptidase inhibitor activity.",
     "L1 elements, as was previously shown for LINEs, were connected with olfactory "
     "receptors (FDR = 6.4 × 10⁻⁴¹), negative regulation of fatty acid "
     "metabolism (FDR = 4.0 × 10⁻⁴) and flavone metabolism "
     "(FDR = 0.031); L2 had non-informative protein binding; and CR1 elements had "
     "surprising connections with neurotransmitter-related processes, retaining synapse "
     "(FDR = 6.9 × 10⁻⁴) and response to cocaine (FDR = 0.027) but losing "
     "regulation of postsynaptic neurotransmitter receptor internalization, ventricular "
     "septum development and serine-type endopeptidase inhibitor activity (all "
     "FDR = 0.086) at the tightened threshold. Flavone metabolic process (GO:0051552, five "
     "overlapping UDP-glucuronosyltransferase genes: UGT1A6, UGT1A7, UGT1A8, UGT1A9 and "
     "UGT1A10) is retained at the L1 family level but not at the LINE class level, where its "
     "FDR is 0.088."),
])
edit(body, "F", "LTR families demonstrated 44 unique GO terms", [
    ("LTR families demonstrated 44 unique GO terms compared to 25 ones in LTRs as a class "
     "(Figure 5A).",
     "LTR families demonstrated 33 unique GO terms compared to 22 in LTRs as a class "
     "(Figure 5A)."),
    ("ERVL elements were connected with bitter taste receptors and with keratins",
     "ERVL elements were connected with bitter taste receptor activity (FDR = 0.023) and "
     "with hair cycle (FDR = 0.005), the keratin and intermediate filament terms falling "
     "just outside the 0.05 threshold (FDR = 0.061 and 0.054)"),
    ("ERV1 elements showed the most diverse set of GO terms and were inserted adjacent to "
     "genes of xenobiotics metabolism, fatty aldehyde dehydrogenases, succinyl-CoA breaking "
     "down, suppression of endosymbionts, arylsulfatase, zinc ion binding, ubiquitin ligase "
     "and sphingolipid metabolism.",
     "ERV1 elements showed the most diverse set of GO terms (16) and were inserted adjacent "
     "to genes of zinc and metal ion binding, transcriptional regulation, lipid and alcohol "
     "metabolism, xenobiotic metabolism, arylsulfatase activity, "
     "N-acylphosphatidylethanolamine metabolism, Wnt signalling and a Cul2-RING ubiquitin "
     "ligase complex."),
])
edit(body, "F", "Alu and MIR insertions-adjacent GO terms were remarkably different", [
    ("MIR elements were inserted in the vicinity of genes of voltage-gated potassium ion "
     "channels, phosphatidylinositol-4,5-bisphosphate binding, arginine deiminases, cellular "
     "response to cadmium, copper and zinc ions, macrophage activation, exosomes, complement, "
     "sensory perception of sound and negative regulation of cytokine signaling.",
     "MIR elements were inserted in the vicinity of genes of voltage-gated potassium ion "
     "channels and potassium transport, protein-arginine deiminases, cellular response to "
     "zinc ion and intracellular zinc ion homeostasis, extracellular exosomes, positive "
     "regulation of the inflammatory response, cytoskeleton and sensory perception of "
     "sound. The metal claim is now zinc-only: cellular response to cadmium ion and "
     "detoxification of copper ion both have FDR = 0.078 and are not retained, as are "
     "macrophage activation (FDR = 0.053), "
     "phosphatidylinositol-4,5-bisphosphate binding, negative regulation of response to "
     "cytokine stimulus (both FDR = 0.062) and the complement component C1q complex "
     "(FDR = 0.078)."),
])
edit(body, "F", "The only two DNA elements that showed non-random enrichment", [
    ("The former was associated with a single significant GO term of cysteine-type "
     "endopeptidase inhibitor activity, the latter related to MHC class I via three GO terms "
     "(Figure 6A).",
     "The former was associated with a single significant GO term of cysteine-type "
     "endopeptidase inhibitor activity (FDR = 0.040), the latter with olfactory receptor "
     "activity (FDR = 0.011) and with MHC class I through a single term, MHC class I protein "
     "complex (FDR = 0.025); beta-2-microglobulin binding and antigen processing and "
     "presentation of peptide antigen via MHC class I both have FDR = 0.051 and are no "
     "longer retained, so the MHC association rests on one term rather than three "
     "(Figure 6A)."),
])
edit(body, "F", "We classified all the family-level GO terms into the main functional", [
    ("Among 22 functional groups (including the “Too general” and “Other” "
     "groups), cell adhesion and cell migration were found only once (in MIR and L1 "
     "elements, respectively). Protein biogenesis processes (protein degradation and "
     "translation) were the next least present group, found in 4 and 3 cases, respectively. "
     "In contrast, too general terms, transcription and nervous system were the most "
     "frequent ones (46, 18 and 13 instances, respectively). Among all the 22 functional "
     "groups and 14 families with GO terms, there was no co-clustering by large-scale "
     "functional metagroups (Figure 6B) and only the RNA processing group was significantly "
     "overrepresented in Alu elements according to the FDR-adjusted Fisher exact test of a "
     "process enrichment in any given family versus all the rest ones.",
     "Among 21 functional groups (including the “Too general” and “Other” "
     "groups), cell migration, cytoskeleton and ribosome biogenesis were found only once "
     "each, and cell adhesion is no longer represented at all. Protein ubiquitination and "
     "degradation was the next least present group, found in 2 cases. In contrast, too "
     "general terms, transcription and its regulation, and the nervous system and sensory "
     "system groups were the most frequent (41, 15 and 9 instances each, respectively). "
     "Among all 21 functional groups and 13 families with GO terms, there was no "
     "co-clustering by large-scale functional metagroups (Figure 6B) and only the RNA "
     "splicing and processing group was significantly overrepresented, in Alu elements "
     "(FDR = 0.027), according to the FDR-adjusted Fisher exact test of a process enrichment "
     "in any given family versus all the rest."),
])
edit(body, "F", "We also compared family-level GO terms with the respective class-level", [
    ("There were 26 functional groups extracted in this analysis (Supplementary File 8).",
     "There were 24 functional groups extracted in this analysis (Supplementary File 8)."),
    ("Both RNA processing and DNA repair terms were enriched for Alu elements, whereas "
     "metals metabolism and other (specific) terms were enriched for MIR TEs. ERV1 elements "
     "had statistically significantly overrepresented lipids and other metabolism groups, "
     "whereas ERVL-MaLR elements had general terms enriched, and ERVK endogenous retroviruses "
     "are also co-localized with lipid metabolism genes as in the case of ERV1 elements.",
     "Both RNA splicing and processing (FDR = 8.7 × 10⁻⁴) and DNA repair "
     "(FDR = 0.002) terms were enriched for Alu elements, whereas other (specific) terms "
     "were enriched for MIR TEs (FDR = 0.013); the metals metabolism association of MIR "
     "elements, significant at the previous threshold (FDR = 0.045), is not retained at "
     "0.05. ERV1 elements had a statistically significantly overrepresented other "
     "metabolism group (FDR = 0.011) but no longer a significant lipids metabolism group "
     "(FDR = 0.045 at 0.1), whereas ERVL-MaLR elements had general terms enriched "
     "(FDR = 0.006) and ERVK endogenous retroviruses remain co-localised with lipid "
     "metabolism genes (FDR = 0.029)."),
])
edit(body, "F", "We than sought to elucidate the main factors determining the TE families", [
    ("We found no significant correlation with enrichment level of TEs in the genes TSS "
     "vicinity (Figure 7A), but TE families whose enrichment near TSS was significant "
     "against the randomized background check showed also higher number of GO terms "
     "(Figure 7B). Expectedly, GO terms number strongly depended on total number of elements "
     "in a family (Figure 7C, Pearson r = 0.633, p-value = 0.015). Additionally, there was no "
     "significant correlation of GO terms count with average divergence in a subfamily "
     "(Figure 7D, 7E), suggesting that TE families of different evolutionary ages can "
     "non-randomly associate with defined functional groups of human genes. Finally, TE "
     "families that showed non-zero number of GO terms had significantly higher enrichment "
     "level (Figure 7F) and total copy number (Figure 7G).",
     "We found no significant correlation with the enrichment level of TEs in the gene TSS "
     "vicinity (Figure 7A, Pearson R = 0.167, raw p = 0.585), and at the tightened threshold "
     "the difference in GO term count between families with and without significant "
     "enrichment against the randomised background is no longer significant either "
     "(Figure 7B, Mann-Whitney U, raw p = 0.113; medians 12 and 2 for 11 and 2 families). "
     "As expected, the GO term number depended strongly on the total number of elements in a "
     "family (Figure 7C, Pearson R = 0.645, raw p = 0.017). There was no significant "
     "relationship between GO term count and average family divergence (Figure 7D, Pearson "
     "R = -0.283, raw p = 0.348; Figure 7E, Mann-Whitney U, raw p = 0.208), suggesting that "
     "TE families of different evolutionary ages can non-randomly associate with defined "
     "functional groups of human genes. Finally, TE families with a non-zero number of GO "
     "terms had a significantly higher enrichment level (Figure 7F, Mann-Whitney U, raw "
     "p = 0.029; median observed/random OR 1.04 against 0.73) and a significantly higher "
     "total copy number (Figure 7G, raw p = 1.2 × 10⁻⁶)."),
])

# =======================================================================================
# Stage G -- Results, new text
# =======================================================================================
t1 = ifna_tests["T1_divergence_unmatched"]
t2 = ifna_tests["T2_divergence_L1_count_matched"]
t3 = ifna_tests["T3_divergence_gene_density_matched"]
t4 = ifna_tests["T4_subfamily_composition"]

IFNA_PARAGRAPHS = [
    "The interferon alpha domain in detail. Because this locus carries the clearest signal "
    "of recent L1 activity in the analysis, we characterised it directly rather than only "
    f"through the gene sets. The {ifna_qc['domain_length_bp']:,} bp domain "
    f"(chr9:21,150,692-21,370,055) contains {ifna_qc['n_tes_in_window']} annotated TEs, of "
    f"which {ifna_qc['n_l1_in_window']} ({ifna_qc['l1_fraction_of_window_tes']:.0%}) are L1 "
    "elements; the remainder are Alu (33), L2 (15), MIR (13), hAT-Charlie (11), ERV1 (10), "
    "ERVL-MaLR (6) and ten elements of eight further families. The L1 density in the domain "
    f"is {ifna_qc['l1_per_mb_window']:.0f} elements per Mb against "
    f"{ifna_qc['l1_per_mb_genome']:.1f} per Mb genome-wide "
    f"({ifna_qc['n_l1_genome']:,} L1 copies over 3.1 Gb), a "
    f"{ifna_qc['l1_density_ratio']:.2f}-fold excess. The mean divergence of these L1 "
    f"elements is {ifna_qc['mean_l1_divergence_window']:.1f} against a genome-wide L1 mean "
    f"of {ifna_qc['mean_l1_divergence_genome']:.1f} (median "
    f"{ifna_qc['median_l1_divergence_genome']:.0f}). The two quantities are different "
    "measurements and should not be confused: the 95.0-161.7 range quoted above is the "
    "per-gene average divergence of the LINE elements around each interferon TSS, whereas "
    f"{ifna_qc['mean_l1_divergence_window']:.1f} is the mean over the individual L1 elements "
    "in the domain.",

    f"The signal is not carried by a handful of outliers. The "
    f"{ifna_qc['n_l1_in_window']} L1 elements span "
    f"{ifna_qc['n_distinct_l1_subfamilies']} distinct subfamilies, including the young "
    "primate-specific L1PA2-L1PA8, L1PA10, L1PA14, L1PA15, L1P1, L1P3, L1P4e, L1P5, L1PB and "
    "L1PREC2 alongside older mammalian L1M clades. Excluding the single element annotated "
    "with divergence 0 - implausible for its L1P3 clade and possibly a RepeatMasker artefact "
    f"- raises the mean only to {ifna_qc['mean_l1_divergence_window_excl_zero']:.1f}, and "
    "removing the five youngest elements altogether still leaves it at "
    f"{ifna_qc['mean_l1_divergence_window_excl_5_youngest']:.1f}, far below the genome-wide "
    "mean.",

    "Formal testing against matched random windows confirms the deficit. Compared with "
    f"{int(t1['n_null_windows']):,} random 220 kb autosomal windows containing at least one "
    f"L1, the domain's mean L1 divergence is {t1['z_score']:.2f} standard deviations below "
    f"the null mean of {float(t1['null_mean']):.1f} (two-sided empirical p = "
    f"{float(t1['empirical_p_two_sided']):.3f}, raw). Restricting the null to windows with at "
    "least 40 L1 elements, which controls for the possibility that the result is a "
    f"by-product of high local L1 density, strengthens it ({int(t2['n_null_windows']):,} "
    f"windows, z = {t2['z_score']:.2f}, p = "
    f"{float(t2['empirical_p_two_sided']):.3f}); so does restricting it to gene-dense "
    f"windows with at least 10 annotated genes and at least one L1 "
    f"({int(t3['n_null_windows']):,} windows, z = {t3['z_score']:.2f}, p = "
    f"{float(t3['empirical_p_two_sided']):.4f}), where only one null window out of "
    f"{int(t3['n_null_windows']):,} is as extreme as the observation. Finally, the domain is "
    "enriched for young primate-specific L1 copies relative to the rest of the genome: "
    f"{int(t4['window_young'])} of its {ifna_qc['n_l1_in_window']} L1 elements belong to "
    f"L1HS or L1P* clades against {int(t4['genome_young']):,} of "
    f"{int(t4['genome_young']) + int(t4['genome_old']):,} genome-wide (Fisher exact odds "
    f"ratio {float(t4['observed']):.2f}, raw p = "
    f"{float(t4['empirical_p_two_sided']):.1e}). The null distributions and the subfamily "
    "composition are shown in Figure 8.",
]
insert_paragraphs_after(
    body, "G", "Finally, DNA elements of highest and lowest divergence", IFNA_PARAGRAPHS)

sens = numbers["sensitivity"]
ROBUSTNESS_PARAGRAPH = [
    "Sensitivity to the window size and the gene-set cut-off. Repeating the whole enrichment "
    "analysis at 5 kb and 20 kb leaves the ordering of TE groups essentially unchanged: the "
    "observed-to-random odds ratio correlates between window pairs with Spearman rho of "
    f"{sens['spearman_rho_range'][0]:.2f} to {sens['spearman_rho_range'][1]:.2f} across the "
    "6 classes and 44 families, and a label-shuffling permutation test rejects chance "
    f"concordance in every pair (largest p = {sens['max_concordance_permutation_p']:.3f}). "
    f"{sens['n_significance_flips']} of the 50 tested groups change significance status "
    "between windows, all of them small families near the threshold (RC/Helitron, ERVL, "
    "Gypsy, I-Jockey, MULE-MuDR, PiggyBac, TcMar, TcMar-Mariner, hAT and tRNA-Deu); no class "
    "other than Helitrons flips. The gene sets themselves are stable, with overlap "
    f"coefficients of {sens['geneset_overlap_coefficient_range'][0]:.2f} to "
    f"{sens['geneset_overlap_coefficient_range'][1]:.2f} between windows and every "
    "hypergeometric p-value against chance below "
    f"{sens['geneset_max_hypergeometric_p']:.0e}, and the gene rankings agree with Kendall "
    f"tau between {sens['kendall_tau_range'][0]:.2f} and "
    f"{sens['kendall_tau_range'][1]:.2f}. The weakest cell throughout is the SVA class "
    "between 5 and 20 kb, which is expected for the smallest class: 6,274 elements "
    "genome-wide give sparse per-gene counts and heavy tying. Widening the Gene Ontology "
    "gene sets from the top and bottom 5 % to 10 % preserves a median of "
    f"{sens['percentile_fraction_preserved_median']:.0%} of the significant terms per group "
    f"and adds many more ({sens['percentile_terms_lost_total']} terms lost, "
    f"{sens['percentile_terms_gained_total']} gained overall), as expected from the larger "
    "foreground. Of the nine abstract-level claims tested at both percentiles, eight survive "
    "both; the exception is the hAT-Charlie MHC class I association, which is significant "
    "only in the 5 % set. All comparisons are given in Supplementary Figure S13 and the "
    "accompanying tables.",
]
robustness_elements = insert_paragraphs_after(
    body, "G", "The overall comparison of functional groups enrichment at the level of "
               "classes", ROBUSTNESS_PARAGRAPH)

LU_PARAGRAPH = [
    "Direct comparison with the previous repeat-based gene categorisation. To place these "
    "results next to the earlier repeat-based classification of genes, we tested the overlap "
    "between our top-5 % gene sets and the three published repeat-enriched gene categories "
    "of that study, mapping their mouse genes to human orthologs through MGI homology "
    "classes and restricting both sides to the same universe of human genes with a mouse "
    "ortholog. The counterpart pairs agree strongly and specifically. Our Alu-enriched genes "
    "overlap their SINE-enriched category 2.87-fold above chance (273 shared genes, odds "
    "ratio 4.02, FDR = 4.4 × 10⁻⁶¹), and our SINE class as a whole "
    "2.71-fold (270 genes, odds ratio 3.69, FDR = 6.1 × 10⁻⁵⁵); our "
    "L1-enriched genes overlap their L1-enriched category 2.22-fold (49 genes, odds ratio "
    "2.44, FDR = 6.3 × 10⁻⁷). The cross pairs are correspondingly depleted: "
    "our SINE set against their L1 category is 0.13-fold "
    "(FDR = 7.6 × 10⁻⁹) and our Alu set against their L1 category 0.41-fold "
    "(FDR = 5.8 × 10⁻⁴). One counterpart pair diverges, and we report it "
    "rather than set it aside: our MIR-enriched genes are depleted, not enriched, in their "
    "SINE-enriched category (0.46-fold, FDR = 2.8 × 10⁻⁹). This is "
    "interpretable, because MIR elements are ancient tRNA-derived SINEs shared across "
    "mammals, whereas the SINE content driving a mouse-based categorisation is dominated by "
    "the young rodent-specific B1 and B2 families; the two are the same class but not the "
    "same elements. The full category-by-group matrix with Fisher p-values and Jaccard "
    "indices is provided as a supplementary table.",
]
# Placed immediately after the sensitivity paragraph so the two "how does this compare"
# results sit together at the end of the Results, before the Discussion opens.
if not structural_edit_needed(body, "G", "insert the Lu et al. overlap paragraph",
                             LU_PARAGRAPH[0][:80]):
    pass
elif isinstance(robustness_elements, list) and robustness_elements:
    lu_elements = [ins_paragraph(text, style=style_of(robustness_elements[-1]))
                   for text in LU_PARAGRAPH]
    insert_after(robustness_elements[-1], lu_elements)
    record("G", "insert the Lu et al. overlap paragraph", True)
else:
    record("G", "insert the Lu et al. overlap paragraph", False)

# =======================================================================================
# Stage J -- Discussion restructured into six subsections
#
# The published Discussion is ~3,970 words in five loosely bounded subsections. Reviewer
# major comment 6 asks for it to be substantially shortened and refocused on (1) findings,
# (2) comparison with prior work, (3) limitations and (4) hypotheses for future testing.
# Per decision D7 the cancer subsection is kept in condensed form, and two further
# subsections are added: the proximity-null argument that answers major comment 3 without
# new analysis, and the mechanistic framework that carries the existing schematic.
#
# Material that leaves the main text is not deleted from the record: it is moved into
# Extensive_discussion_260803.docx by 14_build_extensive_discussion.py, which is built from
# the same baseline so the Mendeley citations in the moved passages stay live.
# =======================================================================================
DISCUSSION_HEADINGS_TO_DELETE = [
    "10 kb TSS neighborhood as an optimal window size based on public literature",
    "Enrichment of TE families and classes near human genes TSS",
    "Functional groups enriched and deficient in TE insertions by classes count",
    "Interferon alpha region of low LINEs divergence",
]

# Every paragraph whose material either moves to the Extensive discussion or is replaced by
# one of the new subsections. Matched on a distinctive opening substring.
DISCUSSION_PARAGRAPHS_TO_DELETE = [
    "In the present study we used a co-mapping window of 10 kb, 5 kb upstream",
    "The current analysis validates this approach by highlighting the importance",
    "As a necessary preliminary analysis, we studied degree of enrichment of TE classes",
    "The enrichment of SINEs and SVAs near TSS (1.468- and 1.368-fold enrichment",
    "Our results of SVA elements being enriched near TSS are connected also to the fact",
    "Moreover, SVA elements are the youngest TE class in humans with active evolutionary",
    "DNA transposons (e.g., TcMar, hAT) utilize a “cut-and-paste” transposase",
    "LTR elements, derived from ancient endogenous retroviruses, are also heavily depleted",
    "At the level of TE families, our results reveal a more complex regulatory landscape.",
    "The CR1 family (Chicken repeats, LINE) also shows significant enrichment",
    "Whereas SINEs showed 1.468-fold enrichment near TSS, their major families",
    "Finally, we observed no connection of family enrichment near TSS with their divergence",
    "On average, according to our results, each human gene harbors 15.05 TEs",
    "Genes involved in embryogenesis, transcription, and nervous system development",
    "Conversely, genes enriched with TEs (the “TE top” group) are involved in RNA "
    "splicing",
    "RNA splicing: Alu elements were early shown to profoundly impact RNA splicing",
    "DNA repair: The enrichment of TEs (particularly Alu elements) near DNA repair genes",
    "Olfactory receptors: L1 elements are specifically enriched near olfactory receptor",
    "Other groups: metals metabolism, other (specific) metabolism, cytoskeleton, cell death",
    "At the level of TE classes, the previous landmark study by",
    "At the level of families, Alu elements are of particular interest.",
    "Alus enrichment near genes of RNA processing is corroborated by the recent evidence",
    "The MIR (Mammalian-wide Interspersed Repeat) family, though less numerous than Alus",
    "SVA elements were previously shown to impact host genes via premature termination",
    "Among the TE classes, only LTR families all demonstrated functional associations",
    "DNA transposons, which have been extinct in the human lineage for millions of years, "
    "are mostly functionally neutral",
    "It is important to note than the fact that certain TE families, precisely 30 out of 44",
    "At the level of lowest and highest divergence by TE classes, ancient high-divergence",
    "One of the most biologically significant findings that we can report is",
    "The enrichment of young LINEs in this region suggests an ongoing evolutionary arms "
    "race.",
    "The contemporary anticancer treatment heavily relies on the single molecule based",
    # Superseded by the Principal findings subsection, which now carries the 99.11 % framing
    # and the 25 % TE-derived comparison; keeping both would state them twice.
    "The analysis reported here shows that 99.11% of unique TSS are within 10 kb of at least",
]

PRINCIPAL_FINDINGS = [
    "Principal findings",
    "Three results carry the paper. First, TE proximity to human genes is close to universal: "
    "99.11 % of the 38,704 unique TSS have at least one annotated TE within 10 kb, and only "
    "343 do not. On average each gene carries 15.05 elements in that window (7.82 SINEs, 4.39 "
    "LINEs, 1.49 DNA elements, 1.32 LTRs, 0.03 SVA elements and 0.004 Helitrons), with CIB3 "
    "as the extreme at 42. Set against earlier estimates that roughly a quarter of human "
    "promoters are TE-derived, the shift from 25 % “TE-derived” to 99.11 % "
    "“TE-proximal” marks the difference between a TE supplying the start site and a "
    "TE merely being available as an auxiliary regulator, insulator or chromatin modifier.",
    "Second, once the length bias of the odds ratio is removed by the permutation background, "
    "enrichment near TSS is modest and class-specific rather than general: SINEs and SVA "
    "elements are enriched (1.468- and 1.368-fold over random expectation) while LINEs, LTRs "
    "and DNA elements are depleted (0.877, 0.667 and 0.938). At the family level only 7 of 44 "
    "families are enriched on both tests and 9 are depleted. The direction of these effects "
    "reverses for several groups if the raw odds ratio is used instead, which is the "
    "methodological point the permutation control exists to make.",
    "Third, the functional associations separate cleanly by TE group. Genes with the fewest "
    "TEs are dominated by embryogenesis, transcription and nervous-system terms, consistent "
    "with core pathways being protected from TE-borne regulatory noise; genes with the most "
    "are enriched for RNA splicing, DNA repair, telomere maintenance and apoptotic signalling. "
    "Individual groups add specificity: SVA elements near genes for termination of RNA "
    "polymerase II transcription, Alu near RNA splicing and DNA repair, MIR near zinc "
    "homeostasis and voltage-gated potassium channels, L1 near olfactory receptors and fatty "
    "acid metabolism, LTR families near ion transport. The interferon-alpha domain on "
    "chromosome 9 is the strongest single locus: 77 of its 175 TEs are L1 copies spanning 36 "
    "subfamilies, at a mean divergence of 135.7 against 188.2 genome-wide, and the deficit "
    "holds against null windows matched for L1 count and for gene density.",
]

COMPARISON_WITH_PRIOR_WORK = [
    "Comparison with prior work",
    "The previous landmark repeat-based categorisation of genes used 20 kb and 2 kb windows "
    "and a region-binning design, profiling TE abundance separately in promoter, intron, "
    "downstream, 5′ UTR, CDS and 3′ UTR compartments, then quantile-normalising and "
    "clustering. It reported SINE-enriched genes associated with the ribosome, translation, "
    "RNA processing, the nucleolus and protein transport, and LINE-enriched genes with "
    "olfactory receptors, retinol metabolism, the epoxygenase P450 pathway and immunoglobulin "
    "domains. Our design is simpler and per-element rather than per-compartment, so the two "
    "are not interchangeable; the gene-set overlap reported in the Results is the direct "
    "comparison the shared quantities allow, and it is strong and specific where the "
    "counterpart categories match (Alu against SINE-enriched, 2.87-fold; L1 against "
    "L1-enriched, 2.22-fold) and correspondingly depleted where they cross.",
    "Whether the residual differences are methodological or a consequence of the assembly "
    "cannot be settled without re-running their pipeline on hg38, which was beyond the scope "
    "of this revision. The assembly-attributable component can, however, be bounded. Taking "
    "the UCSC hs1-to-hg38 liftOver chains and defining newly resolved sequence as chm13v2.0 "
    "bases with no chain coverage plus in-chain gaps of at least 1 kb gives 208.8 Mb, 6.70 % "
    "of the assembly. Only 15,381 of the 3,709,429 TEs (0.41 %), 190 of the 38,704 TSS "
    "windows (0.49 %) and 158 of the 28,738 genes (0.55 %) fall in it, and TE-annotated bases "
    "make up 4.7 % of newly resolved sequence against 82.9 % genome-wide, because the newly "
    "resolved fraction is dominated by satellite rather than interspersed repeat. The "
    "interferon-alpha domain contains no newly resolved sequence at all and is fully "
    "alignable to hg38. The assembly therefore cannot account for more than a fraction of a "
    "percent of the comparison, and the remaining differences are attributable to "
    "methodology: TSS-window proximity against region binning, per-element against per-bin "
    "normalisation, and the length-bias correction by permutation.",
    "Where our results confirm prior work they do so on a complete assembly: SINE and Alu "
    "association with RNA processing, L1 with olfactory receptors, and the SVA association "
    "with premature transcription termination that was predicted from the internal "
    "polyadenylation signal of these elements. Where they differ, the reasons are traceable "
    "to design. An early survey that found Alu elements preferentially near metabolism, "
    "transport and signalling genes examined only chromosomes 21 and 22. A proximity study "
    "that placed SVA elements near zinc finger clusters used hg19 and 1 Mb bins rather than "
    "10 kb windows. Our own comparison of MIR-enriched genes with a mouse-derived "
    "SINE-enriched category is depleted rather than enriched, which follows from MIR being an "
    "ancient tRNA-derived SINE shared across mammals while the rodent SINE content is "
    "dominated by young B1 and B2 families.",
]

LIMITATIONS = [
    "Limitations",
    "This is a correlative and static map. It records where annotated elements lie relative "
    "to annotated transcription start sites in a single genome, and it cannot show that any "
    "element regulates any gene. No expression, chromatin or perturbation data enter the "
    "analysis, so every association reported here is a hypothesis about function rather than "
    "evidence of it, and the causal vocabulary of an evolutionary arms race is used in this "
    "paper only when reporting other people's conclusions.",
    "Five specific limitations bound the interpretation. The analysis rests on one assembly "
    "and one annotation: T2T-CHM13v2.0 is a single haploid genome, so the polymorphic and "
    "population-variable insertions that matter most for recent TE activity are absent by "
    "construction, and RepeatMasker calls in newly resolved regions are less well validated "
    "than elsewhere, although those regions carry under half a percent of the elements "
    "analysed. Gene Ontology annotation is itself biased towards well-studied gene families, "
    "which inflates the apparent specificity of terms such as olfactory receptor activity "
    "where large paralogous families are annotated together. Genes differ systematically in "
    "their number of annotated TSS, and because elements were counted per TSS an element in "
    "two windows of the same gene is counted twice, so genes with many isoforms are "
    "over-weighted; the effect is a known property of this design rather than a correctable "
    "artefact of our implementation. The 10 kb window is a literature convention rather than "
    "a measured optimum, though the sensitivity analysis at 5 and 20 kb shows the ordering of "
    "TE groups is robust to it. Finally, no eQTL, expression or tumour data were integrated, "
    "so the connection between TE proximity and transcriptional consequence remains untested "
    "here.",
]

FUTURE_TESTING = [
    "Hypotheses for future testing",
    "The map is most useful where it produces named, testable predictions rather than general "
    "expectations. Three are specific enough to attempt directly.",
    "First, the interferon-alpha domain. If the young L1 copies in this 220 kb region "
    "contribute to interferon gene regulation, then CRISPR interference targeted at "
    "individual L1PA and L1P insertions in the domain should alter the induction kinetics or "
    "the amplitude of IFNA transcription under interferon or double-stranded RNA stimulation, "
    "with an effect restricted to genes whose TSS lie within the same window. The prediction "
    "is falsifiable in a single cell line and the null result, that the elements are passive "
    "passengers in a locus that simply tolerates insertion, is equally informative.",
    "Second, the SVA-termination association. The SSU72L phosphatase cluster on chromosome 11 "
    "carries three SVA_B copies within a 116 kb region and the association is with genes for "
    "termination of RNA polymerase II transcription rather than with transcription generally. "
    "Deleting the SVA_B insertions at that cluster should change read-through or "
    "polyadenylation site usage at the adjacent phosphatase genes if the elements are "
    "functionally recruited, and should not if the co-localisation is incidental. That the "
    "elements belong to SVA_B rather than to the human-specific SVA_F1 subfamily makes this a "
    "test of an older, potentially fixed relationship rather than of ongoing activity.",
    "Third, and most generally, the proximity values reported here are a baseline that "
    "published mark-based enrichments can be conditioned on. Any claim that a TE family is "
    "enriched for a chromatin mark near genes of a given function should be re-expressed as an "
    "enrichment over the positional expectation for that family and that gene class, which "
    "this map supplies for 6 classes and 44 families genome-wide. The prediction is that a "
    "measurable share of published mark-based associations will shrink once normalised this "
    "way, and that those which survive are the candidates worth pursuing mechanistically.",
]

PROXIMITY_NULL = [
    "Proximity as a null model for TE-epigenomic association studies",
    "A large literature reports that some TE family carries some epigenomic mark near genes "
    "of some function, and reads that pattern as evidence of regulatory recruitment or "
    "domestication. Any such claim requires a positional baseline, because a share of the "
    "signal follows from nothing more than certain TE families being physically closer to "
    "certain classes of gene than chance allows. The results reported here supply exactly "
    "that baseline: genome-wide, at telomere-to-telomere resolution, across 6 classes and 44 "
    "families, with the length bias of the odds ratio removed by permutation. Read this way, "
    "the absence of epigenomic data in this study is the point rather than a gap, since the "
    "proximity layer has to be measured on its own before anything can be conditioned on it.",
    "The orthogonal integration that would test functional relevance is our declared next "
    "step and it exists as a companion study: seven epigenomic modalities (CTCF, H3K27ac, "
    "H3K4me1, H3K9me3, H3K36me3, H3K27me3 and H3K4me3) across twelve human cell lines, "
    "mapped onto 6 TE classes, 44 families and 1,122 subfamilies on the same T2T assembly. "
    "The present proximity map is its background correction, and the two are intended to be "
    "used together: the enrichment of a mark on a TE family near a class of gene should be "
    "read against the positional expectation reported here for that same family and gene "
    "class.",
]

MECHANISTIC_FRAMEWORK_INTRO = [
    "Mechanistic framework",
    "The enrichment and depletion values reported above are the net result of several "
    "processes acting in opposite directions, which the schematic in Figure 9 summarises and "
    "which the following four mechanisms describe. None of them is tested by this design; "
    "they are the framework within which the positional pattern becomes interpretable.",
]

MECHANISTIC_FRAMEWORK_CLOSE = [
    "Read together, these mechanisms account for the direction of the effects we measure "
    "without requiring any of them to be dominant. The enrichment of SINEs and SVA elements "
    "near TSS is consistent with their use of the L1 insertion machinery and its preference "
    "for open chromatin during the hypomethylated window of primordial germ cell "
    "development, combined with the low fitness cost of a short insertion. The depletion of "
    "LINEs and of full-length LTR elements is consistent with purifying selection against "
    "large insertions in regulatory domains and with the sequestration of potent viral "
    "promoters into heterochromatin. The retention of ancient, mutationally decayed families "
    "such as L2 and MIR in gene-proximal regions is consistent with constructive neutral "
    "evolution and exaptation, whereby an initially neutral insertion becomes difficult to "
    "remove once host regulation has come to depend on it. Distinguishing these "
    "possibilities requires the perturbation experiments outlined above; the proximity map "
    "constrains which loci are worth perturbing.",
    "A residual observation supports the neutral end of that range. Thirty-one of the 44 "
    "families - Helitrons, 20 DNA families, 4 SINE families, 6 LINE families and no LTR "
    "family - yield no significant functional gene groups at all, which suggests that their "
    "presence near genes is largely a matter of drift and tolerance rather than adaptive "
    "recruitment.",
]

CANCER_CONDENSED = [
    "Recent evidence points to high transpositional activity in cancer, with more than 500 L1 "
    "insertions per tumour reported in bladder cancer, and large-scale TE de-repression is an "
    "emerging hallmark of the disease. The functional groups identified here define the scope "
    "of human molecular processes that distinct TE classes and families could aberrantly "
    "activate when that de-repression occurs, and therefore which gene sets are worth "
    "examining first in tumour cohorts. The wider argument for building compound "
    "TE-derived gene expression signatures, and the biomarker literature it rests on, is "
    "developed in the extended discussion provided as supplementary material.",
]

# The intro paragraph of the Discussion is kept but pointed at the new structure.
edit(body, "J", "In the present paper, we performed an integrated analysis of human TEs", [
    ("The current article is meant to establish a baseline for these studies by the "
     "co-mapping proximity analysis, since most the epigenetics-based studies are relying on "
     "the same proximity principle.",
     "The current article is meant to establish a baseline for these studies by co-mapping "
     "proximity analysis, since most epigenetics-based studies rely on the same proximity "
     "principle. What follows is organised as the principal findings, their comparison with "
     "prior work, the limitations of the design, the specific hypotheses it generates, its "
     "connection with cancer, its use as a null model for epigenomic association studies, "
     "and the mechanistic framework that makes the positional pattern interpretable."),
])

anchor = paragraph_by(body, "In the present paper, we performed an integrated analysis")
if anchor is not None and not structural_edit_needed(
        body, "J", "insert 4 new Discussion subsections",
        PRINCIPAL_FINDINGS[1]):
    anchor = None
    discussion_sections_inserted = False
else:
    discussion_sections_inserted = anchor is not None
if anchor is None:
    if not any(r["edit"].startswith("insert 4 new Discussion") for r in report):
        record("J", "locate the Discussion opening paragraph", False)
else:
    new_sections = []
    for block in (PRINCIPAL_FINDINGS, COMPARISON_WITH_PRIOR_WORK, LIMITATIONS,
                  FUTURE_TESTING):
        new_sections.append(heading(block[0], 2))
        new_sections += [ins_paragraph(text) for text in block[1:]]
    insert_after(anchor, new_sections)
    record("J", f"insert 4 new Discussion subsections ({len(new_sections)} paragraphs)", True)

    # The cancer subsection keeps its heading, and its first paragraph is condensed in place.
    # Its two citations sit in <w:sdt> controls that split the sentence into three run
    # blocks, so the rewrite is expressed as two block-local edits that leave both controls
    # untouched and therefore still live for Mendeley.
    edit(body, "J", "The recent evidence points out high transpositional activity in cancer", [
        ("The recent evidence points out high transpositional activity in cancer, with more "
         "than 500 L1 insertions per tumor in bladder cancer according to a recent preprint ",
         "Recent evidence points to high transpositional activity in cancer, with more than "
         "500 L1 insertions per tumour reported in bladder cancer "),
        (". Moreover, large-scale TE de-repression is an emerging hallmark of cancer. Our "
         "current work defines the scope of human molecular processes that can be aberrantly "
         "activated in cancer by distinct TE classes and families and potentially targeted by "
         "anticancer therapy ",
         ". Large-scale TE de-repression is an emerging hallmark of the disease, and the "
         "functional groups identified here define the scope of human molecular processes "
         "that distinct TE classes and families could aberrantly activate when that "
         "de-repression occurs - and therefore which gene sets are worth examining first in "
         "tumour cohorts "),
    ])
    insert_paragraphs_after(
        body, "J", "Recent evidence points to high transpositional activity in cancer",
        ["The wider argument for building compound TE-derived gene expression signatures, and "
         "the biomarker literature it rests on, is developed in the extended discussion "
         "provided as supplementary material."])

    # The proximity-null subsection and the mechanistic framework follow the cancer one, so
    # the mechanism list and Figure 9 stay with their own subsection at the end.
    cancer = paragraph_by(body, "Connection of TE enrichment with cancer")
    if cancer is None:
        record("J", "locate the cancer subsection heading", False)
    else:
        tail = paragraph_by(body, "Recent evidence points to high transpositional activity") \
            or paragraph_by(body, "The recent evidence points out high transpositional")
        if structural_edit_needed(body, "J", "insert the proximity-null subsection",
                                  PROXIMITY_NULL[1]):
            proximity = [heading(PROXIMITY_NULL[0], 2)]
            proximity += [ins_paragraph(text) for text in PROXIMITY_NULL[1:]]
            insert_after(tail if tail is not None else cancer, proximity)
            record("J", "insert the proximity-null subsection", True)

    # The four-mechanism list is retained; it gets its own heading and closing paragraphs.
    mechanism_lead = paragraph_by(
        body, "The degree of enrichment or deficiency of TE groups near human genes can be")
    if mechanism_lead is None:
        record("J", "locate the mechanism list lead-in", False)
    else:
        if structural_edit_needed(body, "J", "insert the mechanistic framework intro",
                                  MECHANISTIC_FRAMEWORK_INTRO[1]):
            insert_before(mechanism_lead, [heading(MECHANISTIC_FRAMEWORK_INTRO[0], 2)]
                          + [ins_paragraph(text)
                             for text in MECHANISTIC_FRAMEWORK_INTRO[1:]])
        tracked_replace_safe(mechanism_lead, [
            ("The degree of enrichment or deficiency of TE groups near human genes can be "
             "determined by the following factors (Figure 8):",
             "The degree of enrichment or deficiency of TE groups near human genes can be "
             "determined by the following factors (Figure 9):")])
        figure_caption = paragraph_by(
            body, "Schematic representation of mechanisms impacting distribution of TEs")
        if figure_caption is not None:
            tracked_replace_safe(figure_caption, [
                ("Figure 8. Schematic representation of mechanisms impacting distribution of "
                 "TEs with respect to TSS.",
                 "Figure 9. Schematic representation of mechanisms impacting distribution of "
                 "TEs with respect to TSS.")])
            if structural_edit_needed(
                    body, "J", "insert the mechanistic framework closing paragraphs",
                    MECHANISTIC_FRAMEWORK_CLOSE[0]):
                insert_after(figure_caption,
                             [ins_paragraph(text) for text in MECHANISTIC_FRAMEWORK_CLOSE])
                record("J", "insert the mechanistic framework closing paragraphs", True)
        else:
            record("J", "locate the Figure 8 caption", False)

    # Renumbering the two schematics: old Figure 8 -> 9 and old Figure 9 -> 10, because the
    # interferon-alpha panel set becomes the new main Figure 8 (WP2).
    ring = paragraph_by(body, "This “Ring of Power” functional network illustrates")
    if ring is not None:
        tracked_replace_safe(ring, [("(Figure 9)", "(Figure 10)")])
    ring_caption = paragraph_by(
        body, "Schematic representation of molecular processes groups that were found as")
    if ring_caption is not None:
        tracked_replace_safe(ring_caption, [
            ("Figure 9. Schematic representation of molecular processes groups",
             "Figure 10. Schematic representation of molecular processes groups")])
    record("J", "renumber the two schematics to Figures 9 and 10",
           ring_caption is not None)

    # Everything that moves out of the main text, deleted as tracked deletions.
    deleted = 0
    for needle in DISCUSSION_HEADINGS_TO_DELETE + DISCUSSION_PARAGRAPHS_TO_DELETE:
        paragraph = paragraph_by(body, needle)
        if paragraph is None:
            record("J", f"delete: {needle[:70]}", False)
            continue
        delete_paragraph_safe(paragraph)
        deleted += 1
    record("J", f"tracked-delete {deleted} superseded Discussion "
                f"paragraphs/headings", deleted > 0)
    insert_after(anchor, [note_paragraph(
        "The passages deleted below are moved, not discarded: the window-size literature "
        "review, the per-class and per-family mechanistic review and the cancer biomarker "
        "material are reproduced in Extensive_discussion_260803.docx, submitted as a "
        "supplementary file. Cite it once from this Discussion when the supplementary "
        "numbering is fixed."), note_paragraph(
        "Subsection order: Principal findings, Comparison with prior work, Limitations and "
        "Hypotheses for future testing are placed first, in the order reviewer major comment "
        "6 asks for. Mechanistic framework then follows in its original position, ahead of "
        "the cancer and proximity-null subsections, because moving it would mean relocating "
        "the Figure 9 image paragraph. Drag those three subsections into the order "
        "Mechanistic framework last if you prefer it; it is a one-step move in Word and "
        "changes no text.")])

# =======================================================================================
# Stage K -- Data availability
# =======================================================================================
# The wrong repository URL is inside a <w:hyperlink>, whose runs a block-local rewrite cannot
# reach, so this one sentence is replaced whole: the original paragraph is tracked-deleted
# (which the helper does handle for hyperlinks) and the corrected one inserted after it. The
# new URL is inserted as plain text; Word will re-linkify it, and no citation control is
# involved in this paragraph.
URL_MARKER = ("with the revision-specific scripts and notebooks under revision_G3/ and a "
              "REPRODUCE.md giving the exact run order")
url_paragraph = paragraph_by(body, "All code used for the comprehensive proximity mapping")
if url_paragraph is None:
    record("K", "locate the repository URL paragraph", False)
elif not structural_edit_needed(
        body, "K", "correct the repository URL (paragraph replaced, hyperlink deleted)",
        URL_MARKER):
    pass
else:
    insert_after(url_paragraph, [ins_paragraph(
        "All code used for the comprehensive proximity mapping, statistical analysis and GO "
        "functional networking is available in the GitHub repository "
        "https://github.com/Nikit357/T2T_transposons_genes, with the revision-specific "
        "scripts and notebooks under revision_G3/ and a REPRODUCE.md giving the exact run "
        "order and resolved paths.", style=style_of(url_paragraph))])
    delete_paragraph_safe(url_paragraph)
    record("K", "correct the repository URL (paragraph replaced, hyperlink deleted)", True)
edit(body, "K", "Tables of TE-gene intersections, including divergence and family-level", [
    ("Tables of TE-gene intersections, including divergence and family-level enrichment "
     "statistics, have been deposited in the same GitHub repository.",
     "Tables of TE-gene intersections, including divergence and family-level enrichment "
     "statistics, are deposited in the same repository. Specifically: TEs_on_genes.csv and "
     "TEs_on_genes_counts_subfamilies.csv hold the per-gene intersections; "
     "enrichment_families_with_random.csv and enrichment_subfamilies_with_random.csv hold "
     "the class, family and subfamily enrichment statistics with both raw and FDR-adjusted "
     "p-values; TableS_class_enrichment_full.csv holds the unadjusted Fisher p-values "
     "omitted from Table 2; the GO_tables directory holds the Gene Ontology results with "
     "both P-value and FDR columns; and revision_G3/output/ holds the window and percentile "
     "sensitivity tables, the interferon-alpha domain test results and the gene-set overlap "
     "matrix. The 500-permutation background is published as a compacted per-seed count "
     "store with a MANIFEST.json and the 01c_expand_counts.py reconstructor, so the random "
     "background is reproducible without re-running bedtools shuffle. The Gene Ontology "
     "annotation file goa_human.gaf is not redistributed because of its size; download "
     "instructions are in REPRODUCE.md."),
])
insert_paragraphs_after(
    body, "K", "Tables of TE-gene intersections, including divergence and family-level",
    ["The TE annotation, the TSS 10 kb windows, the gene sets used for Gene Ontology "
     "analysis and the interferon-alpha domain are also available as a UCSC Genome Browser "
     "track hub on the hs1 (T2T-CHM13v2.0) assembly, coloured by TE class to match the "
     "figures. Hub URL: https://nikit357.github.io/T2T_transposons_genes/trackhub/hub.txt. "
     "The domain discussed above opens directly at "
     "https://genome.ucsc.edu/cgi-bin/hgTracks?db=hs1&hubUrl=https://nikit357.github.io/"
     "T2T_transposons_genes/trackhub/hub.txt&position=chr9:21150692-21370055.",
     "An archival snapshot of the repository at the time of publication is deposited in "
     "Zenodo under DOI [ZENODO DOI]."])
ZENODO_NOTE = ("Replace [ZENODO DOI] with the minted Zenodo DOI before submission, and "
               "confirm the track hub URL resolves once the gh-pages branch is published.")
paragraph = paragraph_by(body, "An archival snapshot of the repository")
if paragraph is not None and structural_edit_needed(
        body, "K", "Zenodo DOI placeholder note", ZENODO_NOTE):
    insert_after(paragraph, [note_paragraph(ZENODO_NOTE)])
    record("K", "Zenodo DOI placeholder note", True)

# =======================================================================================
# Stage M -- the supplementary package: five workbooks, six new captions, 27 citations
# =======================================================================================
# Ordering inside this stage is load-bearing:
#   M1 first  -- it tracked-deletes the eight File Sn caption paragraphs, which removes
#                eight of the 27 old labels before anything scans for them;
#   M8 next   -- it renumbers the eleven body citations, each scoped to its own paragraph;
#   M7 last of the three -- it rewrites the Data availability list, whose NEW text contains
#                "File S1".."File S5" and must therefore not be visible to M8.
# The rest are independent.

# Which paragraph cites which old label. The label -> workbook/sheet mapping itself lives in
# FILE_S_LABEL_MAP above, so the renumbering has one source of truth and already_present() can
# recognise text that has already been renumbered. The anchor identifies the paragraph.
FILE_S_CITATION_SITES = [
    ("The complete set of TSS 10 kb neighborhoods used in the mapping", ["File S1"]),
    ("Lists of TEs mapped on each gene TSS, their divergence and classification", ["File S1"]),
    ("whereas LINEs and LTRs were less depleted or even enriched near TSS", ["File S2"]),
    ("There were 1436 genes for each of the major classes", ["File S3"]),
    ("significant terms were extracted (FDR threshold of 0.05 was applied", ["File S4"]),
    ("genes of top and bottom divergence for all TEs", ["File S5"]),
    # One paragraph, two citations: the divergence GO table and the per-TSS element table.
    ("We visualized top 30 GO terms of each TE-divergence groups", ["File S6", "File S1"]),
    ("whereas for the less numerous and/or enriched families we took all", ["File S7"]),
    ("2 DNA families (hAT-Charlie and hAT-Tip100)", ["File S8"]),
    ("functional groups extracted in this analysis", ["File S8"]),
]
FILE_S_CITATIONS = [
    (anchor, [(label, FILE_S_LABEL_MAP[label]) for label in labels])
    for anchor, labels in FILE_S_CITATION_SITES
]

# ---------------------------------------------------------------------------------------
# M1 -- the File Sn captions become five workbook descriptions
# ---------------------------------------------------------------------------------------
FILE_S_CAPTION_ANCHORS = [
    "Genomic coordinates of human TSSs and associated TEs",
    "Enrichment statistics of TE subfamilies within 10 kb regions",
    "Genes exhibiting enrichment or depletion of TEs in their vicinity",
    "GO terms, associated genes, and functional group classifications based on TE enrichment",
    "Genes enriched in TE classes and in all TEs, stratified by high and low divergence",
    "GO terms, associated genes, and functional group classifications for TE groups stratified",
    "Genes enriched in specific TE families based on TE counts",
    "GO terms, associated genes, and functional group classifications for TE family-level",
]


def _sheet_rows(workbook):
    """`sheet (n rows)` for every sheet of a workbook except its README."""
    sheets = supplementary[workbook]["sheets"]
    return ", ".join(f"{name}" for name in sheets if name != "README")


WORKBOOK_CAPTIONS = [
    ("File S1. The transposable element to transcription start site map, and the enrichment "
     "statistics. Sheet TSS_TE_intersections gives all 38,704 TSS neighbourhoods with the "
     "elements within 10 kb of each: counts and mean divergence per class, and the "
     "comma-joined lists of subfamilies, families, classes and divergence values. Counts are "
     "per TSS rather than per gene, so a gene with several annotated TSS contributes several "
     "rows. Sheets enrichment_classes, enrichment_families and enrichment_subfamilies give "
     "the observed and random odds ratios, the raw and FDR-adjusted Fisher p-values and the "
     "empirical permutation p-values for the six classes, the 44 families and all 1,143 "
     "subfamilies respectively."),
    ("File S2. The foreground gene sets used for Gene Ontology analysis, in long format. "
     "Sheet by_TE_group gives the top 5 % of genes by element count for each TE group "
     "together with the TE-top and TE-bottom sets; sheet by_divergence gives the highest- "
     "and lowest-divergence 5 % of genes per class; sheet by_family gives the top 5 % of "
     "genes for each of the 44 families. Each row is one gene with its rank in its set."),
    ("File S3. Gene Ontology enrichment at FDR < 0.05, with the manual functional-group "
     "classification. Sheet GO_TE_groups covers the TE class, TE-top and TE-bottom gene "
     "sets, sheet GO_by_divergence the divergence-stratified sets and sheet GO_by_family the "
     "44 families. Every sheet carries both the raw enrichment p-value and the FDR-adjusted "
     "value, together with the fold enrichment, the overlapping genes and the number of "
     "human genes annotated to the term."),
    ("File S4. The interferon-alpha domain, the assembly bound and the overlap with prior "
     "work. Sheets IFNA_elements, IFNA_tests and IFNA_subfamily_composition give all 175 "
     "elements of the 220 kb domain, the four domain tests with their null distributions and "
     "empirical p-values, and the L1 subfamily composition. Sheet assembly_bound gives the "
     "newly resolved telomere-to-telomere sequence and the share of elements, windows and "
     "genes it contributes, by class, family and chromosome. Sheets "
     "prior_work_overlap_matrix, prior_work_categories and prior_work_shared_genes give the "
     "gene-set comparison with the published mouse categories, mapped through orthology."),
    ("File S5. Sensitivity and robustness analyses. Sheets window_classes, window_families, "
     "window_concordance and window_flips give the enrichment statistics at 5, 10 and 20 kb "
     "with their pairwise concordance and every change of significance status. Sheets "
     "percentile_summary and percentile_terms give the Gene Ontology comparison between the "
     "5 % and 10 % gene-set cut-offs. Sheets GO_grid_index, GO_grid_preservation, "
     "GO_grid_terms, GO_grid_concordance and headline_by_condition give the Gene Ontology "
     "results for all six combinations of window size and percentile, the fraction of "
     "published terms preserved in each, every term gained or lost, and each headline claim "
     "under every condition. Sheets geneset_stability and rank_stability give the gene-set "
     "overlap and rank agreement between window sizes."),
]

# The marker check comes first: a tracked-deleted paragraph is still present in the XML and
# `text_of` still reads its delText, so the old captions stay findable after M1 has run.
if paragraph_by(body, WORKBOOK_CAPTIONS[0][:80]) is not None:
    record("M", "M1 five workbook captions replace the eight File Sn captions", True,
           "already present")
    found_captions = []
else:
    caption_paragraphs = [paragraph_by(body, anchor) for anchor in FILE_S_CAPTION_ANCHORS]
    found_captions = [p for p in caption_paragraphs if p is not None]
if not found_captions:
    pass
elif len(found_captions) != len(FILE_S_CAPTION_ANCHORS):
    record("M", f"M1 found only {len(found_captions)} of "
                f"{len(FILE_S_CAPTION_ANCHORS)} File Sn captions", False)
else:
    insert_after(found_captions[-1],
                 [caption_like(text, found_captions[-1]) for text in WORKBOOK_CAPTIONS])
    for paragraph in found_captions:
        delete_paragraph_fully(paragraph)
    record("M", f"M1 replace {len(found_captions)} File Sn captions with "
                f"{len(WORKBOOK_CAPTIONS)} workbook captions", True)

# ---------------------------------------------------------------------------------------
# M8 -- the eleven body citations (the other sixteen are handled by M1 and M7)
# ---------------------------------------------------------------------------------------
for anchor, pairs in FILE_S_CITATIONS:
    renumber_citations_in(body, "M", anchor, pairs)

# ---------------------------------------------------------------------------------------
# M7 -- Data availability
# ---------------------------------------------------------------------------------------
edit_insertion(body, "M", "TableS_class_enrichment_full.csv holds the unadjusted Fisher", [
    ("TableS_class_enrichment_full.csv holds the unadjusted Fisher p-values omitted from "
     "Table 2",
     "File S1, sheet enrichment_classes holds the unadjusted Fisher p-values omitted from "
     "Table 2"),
])

SUPPLEMENTARY_ITEMS_OLD = (
    "The supplementary items are as follows. File S1: genomic coordinates of the 38,704 "
    "human TSS neighbourhoods with their overlapping TE classes, families, subfamilies and "
    "divergence values. File S2: enrichment statistics for TE subfamilies within the 10 kb "
    "TSS regions, with raw and FDR-adjusted p-values. File S3: genes enriched or depleted in "
    "TEs by major TE group. File S4: GO terms, associated genes and functional-group "
    "classifications for the TE class, TE-top and TE-bottom gene sets. File S5: genes "
    "enriched in TE classes stratified by high and low divergence. File S6: GO terms for the "
    "divergence-stratified groups. File S7: genes enriched in specific TE families. File S8: "
    "GO terms for the family-level analysis. Files S2, S4, S6 and S8 carry both raw and "
    "FDR-adjusted p-values, so any reader can apply an alternative multiple-testing "
    "correction."
)
SUPPLEMENTARY_ITEMS_NEW = (
    "The supplementary tables are provided as five thematic workbooks, each opening with a "
    "README sheet that describes its contents and names the script that produced it. "
    "File S1: the TE-to-TSS map and the enrichment statistics for classes, families and "
    "subfamilies. File S2: the foreground gene sets used for Gene Ontology analysis, in long "
    "format. File S3: the Gene Ontology results at FDR < 0.05 with their functional-group "
    "classification. File S4: the interferon-alpha domain, the newly resolved sequence and "
    "the overlap with prior work. File S5: the window-size, percentile and Gene Ontology "
    "grid sensitivity analyses. The enrichment sheets of File S1 and every Gene Ontology "
    "sheet of File S3 carry both raw and FDR-adjusted p-values, so any reader can apply an "
    "alternative multiple-testing correction."
)
edit_insertion(body, "M", "The supplementary items are as follows",
               [(SUPPLEMENTARY_ITEMS_OLD, SUPPLEMENTARY_ITEMS_NEW)])

# ---------------------------------------------------------------------------------------
# M6 -- the editor note on the Table 1 / Table 2 split
# ---------------------------------------------------------------------------------------
edit_insertion(body, "M", "Tables 1 and 2 replace the original 11-column Table 1", [
    ("the unadjusted Fisher p-values moved to TableS_class_enrichment_full.csv in the "
     "supplementary material",
     "the unadjusted Fisher p-values moved to File S1, sheet enrichment_classes"),
])

# ---------------------------------------------------------------------------------------
# M10 -- Figure 6A's term count, read off network_qc.json rather than remembered
# ---------------------------------------------------------------------------------------
FIG6A_TOP_N = qc_top_n("Fig6A_simplified")
NUMBER_WORD = {4: "four", 5: "five", 6: "six", 7: "seven", 8: "eight", 9: "nine", 10: "ten"}
edit_insertion(body, "M", "Up to nine terms by enrichment p-value were selected", [
    ("Up to nine terms by enrichment p-value were selected for each family",
     f"Up to {NUMBER_WORD[FIG6A_TOP_N]} terms by enrichment p-value were selected for each "
     f"family"),
])

# ---------------------------------------------------------------------------------------
# M9 -- Figure 7's caption carries the statistics the half-size panels no longer show
# ---------------------------------------------------------------------------------------
# Anchored on the sentence TAIL, not the whole sentence: a replacement whose target is the
# entire sentence would show the reviewer the sentence struck through and re-added in full,
# when all that changes is an addition after it.
FIGURE7_STATS_OLD = "the Fisher exact stars in (H) are FDR-corrected."
FIGURE7_STATS_NEW = (
    "the Fisher exact stars in (H) are FDR-corrected. Panels (B), (E), (F) and (G) are drawn "
    "without in-panel annotation, so their group sizes and test results are given here: "
    f"(B) {mwu_sentence('7B', 'significantly enriched families', 'non-significant')}; "
    f"(E) {mwu_sentence('7E', 'families with a significant GO term', 'without')}; "
    f"(F) {mwu_sentence('7F', 'families with a significant GO term', 'without')}; "
    f"(G) {mwu_sentence('7G', 'families with a significant GO term', 'without')}."
)
edit_insertion(body, "M", "Pearson correlations in (A), (C) and (D)",
               [(FIGURE7_STATS_OLD, FIGURE7_STATS_NEW)])
assert FIGURE7_STATS_OLD in FIGURE7_STATS_NEW, (
    "the append-in-place form is intended here; replace_inside_insertions advances past each "
    "replacement so this cannot loop")

# ---------------------------------------------------------------------------------------
# M3 -- Supplementary Figure S8's panel C, and two wrong cross-references
# ---------------------------------------------------------------------------------------
# The published text points at "Supplementary Figure 5C" and "5B" for the Sankey and the
# clustermap. Supplementary Figure S5 has only panels A and B and shows TSS distributions;
# the Sankey and the clustermap are panels C and B of Supplementary Figure S8. The references
# are corrected inside the runs the G2 rename already inserted, so the tracked document shows
# the old label struck through and the corrected one added.
edit_insertion(body, "M", "The overall comparison of functional groups enrichment at the level",
               [("Figure S5C", "Figure S8C"), ("Figure S5B", "Figure S8B")])

# "and TE families." with the full stop is unique to panel C: panel B's mention reads
# "and TE families with significant GO terms".
S8_PANEL_C_OLD = "and TE families."
S8_PANEL_C_NEW = (
    S8_PANEL_C_OLD
    + " Every ribbon is drawn here, whereas Figure 7H shows the same comparison with ribbons "
      "carrying fewer than five GO terms omitted, as stated in its legend; that filter hides "
      f"{sankey_filter['Fig7H']['class_to_group']} class-to-group and "
      f"{sankey_filter['Fig7H']['group_to_family']} group-to-family ribbons, together "
      f"accounting for {sankey_filter['Fig7H']['terms_hidden']} GO terms. Because the filter "
      "applies to the visualisation only, the bar heights in Figure 7H are unfiltered, so its "
      "retained ribbons do not fill the bars they connect."
)
edit(body, "M", "(C) Sankey plot of GO terms count comparison",
     [(S8_PANEL_C_OLD, S8_PANEL_C_NEW)])

# ---------------------------------------------------------------------------------------
# M5 -- Methods: the sensitivity analysis is a grid, not two independent arms
# ---------------------------------------------------------------------------------------
METHODS_SENSITIVITY_OLD = (
    "The Gene Ontology analysis was repeated with the top and bottom 10 % of genes (2,872 "
    "genes) alongside the 5 % sets (1,436 genes) used in the main analysis, at the same FDR "
    "threshold."
)
METHODS_SENSITIVITY_NEW = (
    "The Gene Ontology analysis was repeated with the top and bottom 10 % of genes (2,872 "
    "genes) alongside the 5 % sets (1,436 genes) used in the main analysis, at the same FDR "
    "threshold, and at each of the three window sizes rather than at 10 kb alone, giving six "
    "conditions in total for each of the three levels of analysis. The gene sets for every "
    "condition were built by the same code as the published sets, with only the window or the "
    "percentile changed, so that a difference between conditions cannot be an artefact of two "
    "separate implementations. The number of genes in the background is 28,738 at every window "
    "size, so the 5 % and 10 % cut-offs denote the same fraction throughout. No permutations "
    "were repeated for this grid: it compares Gene Ontology results, and the enrichment odds "
    "ratios reported in Table 1 and Figure 2 are unchanged by it."
)
edit_insertion(body, "M", "The Gene Ontology analysis was repeated with the top and bottom",
               [(METHODS_SENSITIVITY_OLD, METHODS_SENSITIVITY_NEW)])

# ---------------------------------------------------------------------------------------
# M4 -- Results: name the figures and files, and report the grid
# ---------------------------------------------------------------------------------------
_preserved_min = go_grid["min_fraction_published_preserved"]
_preserved_median = go_grid["median_fraction_published_preserved"]
_robust = go_grid["headline_claims_robust_all_six"]
_claims = go_grid["headline_claims_total"]
_rho = go_grid["min_spearman_rho"]
_perm_p = go_grid["max_concordance_permutation_p"]

RESULTS_SENSITIVITY_OLD = "All comparisons are given in Figure S13 and the accompanying tables."
RESULTS_SENSITIVITY_NEW = (
    "Repeating the Gene Ontology analysis at every combination of the three window sizes and "
    "the two gene-set cut-offs shows that the window matters more to it than the cut-off "
    "does. Widening the cut-off always finds more terms, as the larger foreground should, but "
    "widening the window does not, and not even in the same direction at every level: the "
    "number of significant terms falls with window width for the class sets defined by "
    "element count, peaks at 10 kb for the divergence-stratified sets, and rises for the "
    "families. Of the terms reported here, the fraction that remains significant under "
    f"another condition is {_preserved_min:.2f} at worst and {_preserved_median:.2f} at "
    f"median, against {_rho:.2f} as the lowest Spearman correlation of per-group term counts "
    f"against the published condition (every label-shuffling permutation test p <= "
    f"{_perm_p:.3f}); {_robust} of the {_claims} claims tested survive all six conditions, "
    "and all of them hold in the published one. Because the grid varies the gene sets and not "
    "the permutation background, this is a property of which genes enter each set rather than "
    "of the enrichment statistics, which are unchanged. All comparisons are given in "
    "Figures S12 and S13 and in File S5."
)
edit_insertion(body, "M", "All comparisons are given in Figure S13",
               [(RESULTS_SENSITIVITY_OLD, RESULTS_SENSITIVITY_NEW)])

# ---------------------------------------------------------------------------------------
# M2 -- the six captions that were never written (Figures S9-S14)
# ---------------------------------------------------------------------------------------
S9_TOP_N = qc_top_n("S9_full")
S10_TOP_N = qc_top_n("S10_full")
S11_TOP_N = qc_top_n("S11_full")

NEW_FIGURE_CAPTIONS = [
    (f"Figure S9. Complete connection map of GO terms enriched in the gene sets defined by "
     f"the number of TEs of each class near the TSS, of which Figure 4A is the simplified "
     f"view. Up to {S9_TOP_N} terms per group are shown against ten in Figure 4A, with edges "
     f"drawn at a Jaccard index of at least 0.1 and terms of at most 1,000 annotated human "
     f"genes retained. Node colour denotes the FDR-corrected enrichment p-value and node size "
     f"the number of genes in the term; group nodes are diamonds coloured by TE class. "
     f"{qc_wrap_sentence('S9_full')}Connection lines are drawn at a capped stroke width for "
     f"legibility; the cap is a drawing width only and the underlying edge weights are "
     f"unchanged. GO terms are counted at FDR < 0.05."),
    (f"Figure S10. Complete connection map of GO terms for the divergence-stratified gene "
     f"sets, of which Figure 5A is the simplified view. Groups are each TE class split into "
     f"its highest- and lowest-divergence tails, together with the same split for all TEs "
     f"combined. Up to {S10_TOP_N} terms per group are shown, with the same edge and term-size "
     f"filters and the same capped line widths as Figure S9. {qc_wrap_sentence('S10_full')}"
     f"GO terms are counted at FDR < 0.05."),
    (f"Figure S11. Complete connection map of GO terms per TE family, of which Figure 6A is "
     f"the simplified view. Up to {S11_TOP_N} terms per family are shown, with the same edge "
     f"and term-size filters and the same capped line widths as Figure S9, and node colour "
     f"denotes the FDR-corrected enrichment p-value. {qc_wrap_sentence('S11_full')}Node size "
     f"and connection-line width follow the same scales as Figure 6A. At this term count the "
     f"label field remains saturated and a few labels still overlap; Figure 6A is the legible "
     f"view of the same network and this panel is provided for the full structure. GO terms "
     f"are counted at FDR < 0.05."),
    ("Figure S12. Robustness of the Gene Ontology results across window size and gene-set "
     "cut-off. (A) Number of significant GO terms in each of the six conditions, three TSS "
     "window sizes by two gene-set percentiles, for each of the three levels of analysis; "
     "cell colour is the term count relative to the published condition, which is outlined. "
     "(B) For each condition, the fraction of the published condition's terms that remain "
     "significant and the Jaccard index of the two term sets. Jaccard is deflated by gains, "
     "so the preserved fraction is the measure that answers whether a published result "
     "survives. (C) Which TE groups keep at least one significant GO term in each condition, "
     "ordered by the number of conditions survived. Terms are counted at FDR < 0.05 "
     "throughout. The grid varies the gene sets only; no permutations were repeated, so the "
     "enrichment odds ratios are unchanged by it."),
    ("Figure S13. Concordance of the enrichment analysis across TSS window sizes and of the "
     "Gene Ontology analysis across gene-set cut-offs. (A) Observed to random odds ratio for "
     "every TE class and family at 5, 10 and 20 kb; hue denotes the TE class and tint the "
     "window size. (B) Bland-Altman comparison of the 44 families between each pair of "
     "windows, with the mean difference and the limits of agreement; points are coloured by "
     "TE class. (C) Number of significant GO terms per group at the 5 % and 10 % cut-offs, "
     "and the fraction of the 5 % terms that remain significant at 10 %; the left annotation "
     "strip carries the TE class colours used in Figures 4B, 5B and 6B. (D) Each "
     "abstract-level claim under all six combinations of window size and cut-off, with filled "
     "circles coloured by the FDR-corrected enrichment p-value on the same scale as the "
     "connection maps, and open circles marking a claim that is not significant or absent. "
     "(E) Overlap coefficient of the top 5 % gene sets and Kendall correlation of the gene "
     "rankings between window pairs, with 95 % bootstrap confidence intervals."),
    ("Figure S14. Convergence of the permutation background at 500 permutations. (A) Running "
     "mean of the random odds ratio for each TE class as permutations accumulate. (B) Running "
     "standard deviation of the same quantity. (C) Drift of the running mean for all 44 "
     "families, expressed in units of the final standard deviation, so that families with "
     "different absolute odds ratios are comparable. By 250 permutations the running mean is "
     "within 0.06 standard deviations of its final value for the worst class and within 0.10 "
     "for the worst family, which is what justifies reporting the background at 500."),
]

s8_caption = paragraph_by(body, "Figure S8. Genes and molecular processes enriched with TE")
if s8_caption is None:
    record("M", "M2 locate the Figure S8 caption", False)
elif not structural_edit_needed(body, "M", "M2 six new captions, Figures S9 to S14",
                               NEW_FIGURE_CAPTIONS[0][:80]):
    pass
else:
    insert_after(s8_caption, [caption_like(text, s8_caption)
                              for text in NEW_FIGURE_CAPTIONS])
    record("M", f"M2 insert {len(NEW_FIGURE_CAPTIONS)} new captions (Figures S9 to S14)", True)


# =======================================================================================
# Report
# =======================================================================================
if not DRY_RUN:
    document.save(TARGET)

outcomes = {"applied": 0, "already present": 0, "not found": 0}
for row in report:
    outcomes[row.get("outcome", "applied" if row["applied"] else "not found")] += 1
failed = [row for row in report if not row["applied"]]

print(f"{len(report)} edits: {outcomes['applied']} applied, "
      f"{outcomes['already present']} already present (skipped), "
      f"{outcomes['not found']} NOT FOUND")
for row in failed:
    print(f"  NOT FOUND [{row['stage']}] {row['edit']}")
with open(os.path.join(HERE, "output", "manuscript_edit_report.json"), "w") as handle:
    json.dump({"target": os.path.basename(TARGET), "in_place": IN_PLACE,
               "dry_run": DRY_RUN, "outcomes": outcomes, "edits": report},
              handle, indent=2)
print("report -> output/manuscript_edit_report.json")
# `not found` is still fatal: an edit whose target text is nowhere in the document means
# either the document is not the one this script was written against, or the edit is wrong.
if failed:
    sys.exit(1)
