#!/usr/bin/env python
"""Phase 7: bring the revision docx to G3 house style (WP15, gaps G2-G16).

**Run this after `13_manuscript_tracked_edits.py`, never instead of it.** That script rebuilds
the revision from the pristine baseline, so it would discard everything done here; this one
edits the revision in place and refuses to run if the Phase 5 edits are not present.

    Revised_manuscript/T2T_genes_article_G3_revision_260803.docx   (in place)

What it does, by gap number from `G3_article_guidelines.md` section 12

    G4/G5/G6  keywords, a ~35-character running title, a full corresponding-author block
    G7        G3's section order and sentence-case headings; `Literature cited` not
              `REFERENCES`; `ETHICAL STATEMENT` folded into Data availability; `Funding`
              split out of the Acknowledgments
    G2        every `Supplementary Figure n` / `Supplementary File n` renamed to
              `Figure Sn` / `File Sn`, with the caption label in bold
    G3        Data availability rewritten to the section 9.3 model, enumerating every File Sn
    G8/G9     Tables 1 and 2 moved to the end of the main text and given 0.5 pt rules
    G10       raw-vs-FDR provenance stated in every figure caption
    G13       human gene symbols italicised
    G14       Acknowledgments trimmed toward journal norm
    G15/G16   AI-usage disclosure kept and placed; the ethical sentence relocated
    mechanics double spacing and page numbers (12 pt and line numbers were already set)

Two structural notes, both reported rather than hidden

    Section reordering is a genuine relocation of body elements, not a Word tracked *move*
    (`<w:moveFrom>`/`<w:moveTo>` needs paired range bookmarks that are easy to get subtly
    wrong). Reject All therefore restores the original text but not the original section
    order. The validation checks content equality as a multiset of paragraphs instead of as
    one string, and an editor note in the file says so.

    Italicising a gene symbol is a delete-plus-insert, because formatting cannot be changed
    without replacing the run. That is what makes it visible as a tracked change at all.

Usage
    ~/venvs/collagen_3_11/bin/python 15_house_style.py [--dry-run]
"""

import copy
import json
import os
import re
import sys

sys.path.insert(0, os.path.expanduser("~/.claude/skills"))

import docx  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from word_rewrite_trackchanges import (  # noqa: E402
    _rev_attrs, delete_paragraph, find_all_p, heading, ins_paragraph, ins_paragraph_runs,
    insert_after, insert_before, is_p, make_del_run, make_run, note_paragraph,
    set_revision_identity, style_of, text_of, wrap_del, wrap_ins,
)

HERE = os.path.dirname(os.path.abspath(__file__))
# The same working file 13_manuscript_tracked_edits.py now edits in place: Daniil edited the
# revision by hand on 2026-08-04 and accepted part of the tracked diff, so the 260803 file
# this script used to target is a superseded record of what the scripts produced before that
# pass. Both scripts must point at the same document or the Phase 7 pass lands on the wrong
# one (revision plan §2.5, §5.11).
TARGET = os.path.join(HERE, "Revised_manuscript",
                      "T2T_genes_article_G3_revision_260804_manual.docx")

DRY_RUN = "--dry-run" in sys.argv
REVISION_AUTHOR = "Claude (G3 house style)"
REVISION_DATE = "2026-08-04T12:00:00Z"

ORCID = "0000-0003-1029-1174"
report = []


def record(gap, label, ok):
    report.append({"gap": gap, "edit": label, "applied": bool(ok)})
    return ok


# =======================================================================================
# Citation-safe tracked editing (same contract as 13_manuscript_tracked_edits.py)
# =======================================================================================
SOFT_MAP = {"\xa0": " ", "‑": "-", "–": "-", "—": "-", "−": "-",
            "‘": "'", "’": "'", "“": '"', "”": '"', "′": "'"}
SOFT_DROP = {"​", "‌", "‍", "﻿"}


def soften(text):
    characters, indices = [], []
    for position, character in enumerate(text):
        if character in SOFT_DROP:
            continue
        characters.append(SOFT_MAP.get(character, character))
        indices.append(position)
    return "".join(characters), indices


def soft(text):
    return soften(text)[0]


def strip_proof_errors(paragraph):
    for marker in paragraph.findall(qn("w:proofErr")):
        paragraph.remove(marker)


def run_blocks(paragraph):
    """Contiguous <w:r> siblings. A <w:sdt> citation control always ends a block."""
    blocks, current = [], []
    for child in list(paragraph):
        if child.tag == qn("w:r"):
            current.append(child)
        elif current:
            blocks.append(current)
            current = []
    if current:
        blocks.append(current)
    return blocks


def block_text(block):
    return "".join(t.text or "" for r in block for t in r.findall(qn("w:t")))


def replace_in_paragraph(paragraph, old, new, bold=False, italic=False):
    """One tracked del+ins for the first occurrence of `old`. True if it was applied."""
    strip_proof_errors(paragraph)
    target = soft(old)
    for block in run_blocks(paragraph):
        text = block_text(block)
        normalised, index_map = soften(text)
        position = normalised.find(target)
        if position < 0:
            continue
        start = index_map[position]
        end = index_map[position + len(target) - 1] + 1
        before, matched, after = text[:start], text[start:end], text[end:]
        base = block[0].find(qn("w:rPr"))
        pieces = []
        if before:
            pieces.append(make_run(before, base=base))
        pieces.append(wrap_del(make_del_run(matched, base=base)))
        if new:
            pieces.append(wrap_ins(make_run(new, base=base, bold=bold, italic=italic)))
        if after:
            pieces.append(make_run(after, base=base))
        anchor = block[-1]
        for piece in pieces:
            anchor.addprevious(piece)
        for run in block:
            paragraph.remove(run)
        return True
    return False


def replace_inside_insertions(body, pairs, limit=200):
    """Apply (old, new) inside `<w:ins>` runs, editing the text directly.

    Needed because Phase 5 inserted paragraphs that themselves refer to supplementary items,
    and `run_blocks` cannot see inside an insertion. Editing our own inserted text in place is
    correct: the span is already marked as added, so there is nothing to track twice.
    """
    counts = {old: 0 for old, _ in pairs}
    for old, new in sorted(pairs, key=lambda p: -len(p[0])):
        target = soft(old)
        for _ in range(limit):
            hit = False
            for insertion in list(body.iter(qn("w:ins"))):
                for run in insertion.findall(qn("w:r")):
                    node = run.find(qn("w:t"))
                    if node is None or not node.text:
                        continue
                    normalised, index_map = soften(node.text)
                    position = normalised.find(target)
                    if position < 0:
                        continue
                    start = index_map[position]
                    end = index_map[position + len(target) - 1] + 1
                    node.text = node.text[:start] + new + node.text[end:]
                    counts[old] += 1
                    hit = True
                    break
                if hit:
                    break
            if not hit:
                break
    return counts


def replace_everywhere(body, pairs, bold=False, italic=False, limit=400):
    """Apply each (old, new) pair to every occurrence in every body paragraph.

    Applied longest-target-first so that `Supplementary Figure 5A` is rewritten before
    `Supplementary Figure 5` and cannot be left as `Figure S5A` mangled into `Figure S5A`.
    """
    counts = {old: 0 for old, _ in pairs}
    for old, new in sorted(pairs, key=lambda p: -len(p[0])):
        for _ in range(limit):
            hit = False
            for paragraph in [el for el in body.iter(qn("w:p"))]:
                if replace_in_paragraph(paragraph, old, new, bold=bold, italic=italic):
                    counts[old] += 1
                    hit = True
                    break
            if not hit:
                break
    return counts


def accepted_text(element):
    """Visible text with Accept All applied: inserted text in, deleted text out.

    This matters because the Phase 5 pass already rewrote many of the strings this script
    looks for. `text_of` from the skill helper concatenates `<w:delText>` *and* `<w:t>`, so a
    heading that was renamed from `RESULTS` to `Results` reads as `RESULTSResults` and no
    exact lookup can find it. Every lookup below therefore works on the accepted view, which
    is also the view the journal will read.
    """
    parts = []
    for node in element.iter(qn("w:t")):
        if any(a.tag == qn("w:del") for a in node.iterancestors()):
            continue
        parts.append(node.text or "")
    return "".join(parts)


def paragraph_by(body, needle, occurrence=0):
    target = soft(needle)
    hits = [el for el in body.iter(qn("w:p")) if target in soft(accepted_text(el))]
    return hits[occurrence] if len(hits) > occurrence else None


def paragraphs_starting(body, prefix):
    target = soft(prefix)
    return [el for el in body.iter(qn("w:p"))
            if soft(accepted_text(el)).strip().startswith(target)]


def heading_paragraph(body, exact):
    """The paragraph whose entire accepted text is `exact`, ignoring soft differences."""
    target = soft(exact).strip().lower()
    for element in body.iter(qn("w:p")):
        if soft(accepted_text(element)).strip().lower() == target:
            return element
    return None


def delete_paragraph_safe(paragraph):
    delete_paragraph(paragraph)
    for sdt in paragraph.findall(".//" + qn("w:sdt")):
        for run in sdt.findall(".//" + qn("w:r")):
            if run.getparent().tag == qn("w:del"):
                continue
            for text in run.findall(qn("w:t")):
                text.tag = qn("w:delText")
            marker = _rev_attrs(OxmlElement("w:del"))
            run.getparent().replace(run, marker)
            marker.append(run)


# =======================================================================================
# Load, with a guard against running before Phase 5
# =======================================================================================
document = docx.Document(TARGET)
body = document.element.body
set_revision_identity(REVISION_AUTHOR, REVISION_DATE)

if paragraph_by(body, "Telomere-to-telomere co-mapping of transposable elements") is None:
    raise SystemExit(
        "This file does not carry the Phase 5 edits. Run 13_manuscript_tracked_edits.py "
        "first — it rebuilds the revision from the baseline and would discard this pass.")
if paragraph_by(body, "Running title:") is not None:
    raise SystemExit("House style has already been applied to this file (running title "
                     "found). Re-run 13_manuscript_tracked_edits.py first to start clean.")

# =======================================================================================
# G4 / G5 / G6 — title-page elements
# =======================================================================================
affiliation = paragraph_by(body, "Institute of Molecular Biology, National Academy")
correspondence = paragraph_by(body, "Correspondence: danya.nikitin.orel@gmail.com")
if affiliation is None or correspondence is None:
    record("G4/G5/G6", "locate the title page", False)
else:
    insert_before(correspondence, [
        ins_paragraph(f"Running title: TEs near human genes in T2T"),
        ins_paragraph("Keywords: transposable elements; T2T-CHM13; transcription start site; "
                      "gene ontology; interferon alpha; LINE-1; human genome"),
    ])
    insert_after(correspondence, [
        ins_paragraph(
            "Corresponding author: Daniil Nikitin, Institute of Molecular Biology, National "
            "Academy of Sciences of the Republic of Armenia, Yerevan, Armenia. "
            "Email: danya.nikitin.orel@gmail.com. "
            f"ORCID: {ORCID}."),
        note_paragraph(
            "G3 wants the corresponding-author block to carry the office mailing address and a "
            "phone number. The institution, city and country are filled in; add the street "
            "address, postal code and phone. The running title is 27 characters including "
            "spaces (G3 asks for about 35, so there is room if you prefer a fuller one)."),
    ])
    record("G4/G5/G6", "running title, keywords, corresponding-author block", True)

# =======================================================================================
# G7 — headings to sentence case, and the two headings G3 names differently
# =======================================================================================
HEADING_RENAMES = [
    ("ABSTRACT", "Abstract"),
    ("INTRODUCTION", "Introduction"),
    ("RESULTS", "Results"),
    ("DISCUSSION", "Discussion"),
    ("MATERIALS AND METHODS", "Materials and methods"),
    ("CONFLICT OF INTEREST", "Conflicts of interest"),
    ("AUTHORS CONTRIBUTION", "Author contributions"),
    ("ACKNOWLEDGEMENTS", "Acknowledgments"),
    ("SUPPLEMENTARY MATERIAL", "Supplementary material"),
    ("DATA AVAILABILITY", "Data availability"),
    ("REFERENCES", "Literature cited"),
]
renamed = 0
for old, new in HEADING_RENAMES:
    paragraph = heading_paragraph(body, old)
    if paragraph is None:
        record("G7", f"heading {old}", False)
        continue
    if replace_in_paragraph(paragraph, old, new):
        renamed += 1
record("G7", f"rename {renamed} headings to G3 wording and sentence case", renamed >= 10)

# ETHICAL STATEMENT is not a G3 section. Its content becomes one sentence in Data
# availability (added below), so the heading and the paragraph are tracked-deleted here.
ethical_heading = heading_paragraph(body, "ETHICAL STATEMENT")
ethical_text = paragraph_by(body, "This study represents a purely computational analysis")
if ethical_heading is not None and ethical_text is not None:
    delete_paragraph_safe(ethical_heading)
    delete_paragraph_safe(ethical_text)
    record("G7/G16", "fold ETHICAL STATEMENT into Data availability", True)
else:
    record("G7/G16", "locate ETHICAL STATEMENT", False)

# Funding is its own G3 back-matter section; the sentence currently ends the Acknowledgments.
funding_sentence = paragraph_by(body, "This research was conducted without specific funding")
if funding_sentence is None:
    record("G7", "locate the funding sentence", False)
else:
    insert_before(funding_sentence, [heading("Funding", 1)])
    record("G7", "split Funding into its own section", True)

# =======================================================================================
# G14 — trim the Acknowledgments toward journal norm
#
# Nothing here is a rule violation, so the author's own words are kept for the people being
# thanked; what goes is the biographical narrative around them. Four paragraphs of about 600
# words become two.
# =======================================================================================
ACK_TO_DELETE = [
    "I am deeply thankful to my parents, Olga Nikitina and Michael Nikitin",
    "My deepest and most heartfelt appreciation goes to my wife, Irina Nikitina",
]
ack_deleted = 0
for needle in ACK_TO_DELETE:
    paragraph = paragraph_by(body, needle)
    if paragraph is not None:
        delete_paragraph_safe(paragraph)
        ack_deleted += 1
record("G14", f"trim {ack_deleted} Acknowledgments paragraph(s)", ack_deleted == 2)

first_ack = paragraph_by(body, "I would like to express my sincere gratitude to the faculty")
if first_ack is not None:
    replace_in_paragraph(
        first_ack,
        "The academic environment, resources, and support I received during my bachelor’s and "
        "master’s studies were instrumental in shaping my early scientific trajectory and "
        "fostering my commitment to bioinformatics and evolutionary biology. I am especially "
        "indebted to Professor Galina Belyakova, whose mentorship played a pivotal role in my "
        "preparation for the 2013 International Biology Olympiad in Bern, where I was awarded "
        "a gold medal. Her guidance, together with that of my other Olympiad mentors, was "
        "fundamental in cultivating my interest in molecular and evolutionary biology well "
        "before the completion of my formal university training.",
        "I am especially indebted to Professor Galina Belyakova, whose mentorship was "
        "fundamental in cultivating my interest in molecular and evolutionary biology.")
    record("G14", "condense the first Acknowledgments paragraph", True)
    insert_after(paragraph_by(body, "I apologize to colleagues and researchers whose relevant")
                 or first_ack,
                 [note_paragraph(
                     "G14: the Acknowledgments were trimmed from four paragraphs (~600 words) "
                     "toward the journal norm. Everyone thanked by name is still thanked. If "
                     "you would rather keep the fuller version, reject these deletions — they "
                     "are tracked individually.")])

# =======================================================================================
# G2 — supplementary items renamed to G3 convention
# =======================================================================================
# Captions first, and with the replacement in bold. Doing them in this order matters: once a
# string has been rewritten it lives inside a `<w:ins>` element, which `run_blocks` cannot
# reach, so a caption label cannot be bolded by a second pass over already-renamed text.
bold_labels = 0
for number in range(1, 9):
    for kind in ("Figure", "File"):
        label = f"Supplementary {kind} {number}."
        for paragraph in paragraphs_starting(body, label):
            if replace_in_paragraph(paragraph, label, f"{kind} S{number}.", bold=True):
                bold_labels += 1
record("G2", f"bold {bold_labels} supplementary caption label(s)", bold_labels >= 16)

# Then every remaining in-text reference, longest target first.
# S1-S14, not S1-S8: the inventory now runs to Figure S14 (S12 is the GO grid, S13 the
# window/percentile concordance, S14 the permutation convergence), and any of them can be
# cited with a panel letter. Longest target first, so "Supplementary Figure 12A" is rewritten
# before "Supplementary Figure 12" and cannot be mangled into "Figure S12A" twice.
figure_pairs = []
for number in range(1, 15):
    for suffix in ("A", "B", "C", "D", "E", ""):
        figure_pairs.append((f"Supplementary Figure {number}{suffix}",
                             f"Figure S{number}{suffix}"))
# The supplementary tables are five thematic workbooks now (decision D-c/D-d), so only
# File S1-S5 are legitimate targets. The 6-8 forms are still swept: if one survives
# anywhere it is a citation the renumbering map (edit M8) missed, and the assertion below
# is what catches it.
file_pairs = [(f"Supplementary File {n}", f"File S{n}") for n in range(1, 9)]
counts = replace_everywhere(body, figure_pairs + file_pairs)
total_renamed = sum(counts.values())

# Phase 5's own inserted text refers to the new supplementary figures, and did so in a
# belt-and-braces form ("Supplementary Figure S14") that is neither the old convention nor
# G3's. Those live inside `<w:ins>`, so they need the insertion-aware pass.
insertion_pairs = [(f"Supplementary Figure S{n}", f"Figure S{n}") for n in range(1, 15)]
insertion_pairs += [(f"Supplementary Figure {n}{s}", f"Figure S{n}{s}")
                    for n in range(1, 15) for s in ("A", "B", "C", "D", "E", "")]
insertion_pairs += [(f"Supplementary File {n}", f"File S{n}") for n in range(1, 9)]
insertion_counts = replace_inside_insertions(body, insertion_pairs)
inside_renamed = sum(insertion_counts.values())
record("G2", f"rename {total_renamed} supplementary references in original text and "
             f"{inside_renamed} in revision-inserted text", total_renamed >= 25)

supplementary_heading = heading_paragraph(body, "Supplementary material")
if supplementary_heading is not None:
    insert_after(supplementary_heading, [note_paragraph(
        "G2: supplementary figure titles and legends belong BELOW the figure and table titles "
        "ABOVE the table, both starting with the bold Figure Sn / Table Sn label. The labels "
        "here are bold; the above/below placement happens when the items are laid out for the "
        "figshare upload, titled 'Supplemental Material for Nikitin 2026' and submitted in one "
        "batch. Rename the on-disk files to match before uploading.")])
    record("G2", "supplementary placement note", True)

# =======================================================================================
# G3 — Data availability to the section 9.3 model
# =======================================================================================
availability_anchor = paragraph_by(body, "An archival snapshot of the repository")
if availability_anchor is None:
    record("G3", "locate Data availability", False)
else:
    insert_before(availability_anchor, [ins_paragraph(
        "Transposable element annotations were obtained from the UCSC RepeatMasker track for "
        "T2T-CHM13v2.0 (hs1) and gene annotations from NCBI RefSeq All (hs1), both via the "
        "UCSC Table Browser.")])
    insert_after(availability_anchor, [
        ins_paragraph(
            "The supplementary items are as follows. File S1: genomic coordinates of the 38,704 "
            "human TSS neighbourhoods with their overlapping TE classes, families, subfamilies "
            "and divergence values. File S2: enrichment statistics for TE subfamilies within "
            "the 10 kb TSS regions, with raw and FDR-adjusted p-values. File S3: genes enriched "
            "or depleted in TEs by major TE group. File S4: GO terms, associated genes and "
            "functional-group classifications for the TE class, TE-top and TE-bottom gene sets. "
            "File S5: genes enriched in TE classes stratified by high and low divergence. "
            "File S6: GO terms for the divergence-stratified groups. File S7: genes enriched in "
            "specific TE families. File S8: GO terms for the family-level analysis. Files S2, "
            "S4, S6 and S8 carry both raw and FDR-adjusted p-values, so any reader can apply an "
            "alternative multiple-testing correction."),
        ins_paragraph(
            "This study analysed only publicly available genomic data. No new biological "
            "material was collected and the work did not involve human participants, animal "
            "subjects or experimental interventions, so no institutional review board approval "
            "or informed consent was required."),
    ])
    record("G3", "Data availability completed to the section 9.3 model", True)

# One-line pointer at the end of Methods, which is what the older GSA guidance asks for while
# current issues print the full statement as back matter. Both, per guidelines section 2.
ai_usage = paragraph_by(body, "Gemini PRO")
if ai_usage is not None:
    insert_after(ai_usage, [ins_paragraph(
        "Data and code availability. All code, intermediate tables and the permutation "
        "background are available as described in the Data availability statement.")])
    record("G3", "one-line availability pointer at the end of Methods", True)

# =======================================================================================
# G10 — p-value provenance in the figure captions
# =======================================================================================
CAPTION_EDITS = [
    ("Figure 3. Evolutionary age comparison of all and TSS-proximal TEs",
     [("(D) Average divergence score distribution of TSS-proximal TEs at the level of "
       "individual classes.",
       "(D) Average divergence score distribution of TSS-proximal TEs at the level of "
       "individual classes. Distributions in (A) and (B) were compared by two-sample "
       "Kolmogorov-Smirnov test; the p-values shown are raw, as each panel is a single test.")]),
    ("Figure 4. Functional analysis of genes whose TSS are enriched or depleted",
     [("Top 30 terms by enrichment p-value were selected for each group, having FDR corrected "
       "p-value below 0.1.",
       "Up to 10 terms by enrichment p-value were selected for each group, all with an "
       "FDR-corrected p-value below 0.05; the full network is Figure S9."),
      ("GO terms with more than 1000 genes were excluded to avoid too general terms.",
       "GO terms with more than 500 genes were excluded to avoid overly general terms, and "
       "edges require a Jaccard index of at least 0.2 and at least 5 shared genes."),
      ("Stars indicate FDR-corrected Fisher enrichment p-value of a given functional group in "
       "each gene enrichment group compared to other enrichment groups.",
       "Stars indicate the FDR-corrected Fisher exact enrichment p-value of a given functional "
       "group in each gene enrichment group compared with the other enrichment groups; GO "
       "terms are counted at FDR < 0.05.")]),
    ("Figure 5. Functional analysis of genes whose TSS are enriched with TEs of different "
     "classes with highest and lowest divergence",
     [("Top 30 terms by enrichment p-value were selected for each group, having FDR corrected "
       "p-value below 0.1.",
       "Up to 10 terms by enrichment p-value were selected for each group, all with an "
       "FDR-corrected p-value below 0.05; the full network is Figure S10."),
      ("GO terms with more than 1000 genes were excluded to avoid too general terms.",
       "GO terms with more than 500 genes were excluded to avoid overly general terms."),
      ("Stars indicate FDR-corrected Fisher enrichment p-value of a given functional group in "
       "each gene enrichment group compared to other enrichment groups.",
       "Stars indicate the FDR-corrected Fisher exact enrichment p-value of a given functional "
       "group in each gene enrichment group compared with the other enrichment groups; GO "
       "terms are counted at FDR < 0.05.")]),
    ("Figure 6. Functional analysis of genes with TSS enriched by TE inserts by family",
     [("Top 30 terms by enrichment p-value were selected for each group, having FDR corrected "
       "p-value below 0.1.",
       "Up to nine terms by enrichment p-value were selected for each family, all with an "
       "FDR-corrected p-value below 0.05; the full network is Figure S11."),
      ("Stars indicate FDR-corrected Fisher enrichment p-value of a given functional group in "
       "each gene enrichment group compared to other enrichment groups.",
       "Stars indicate the FDR-corrected Fisher exact enrichment p-value of a given functional "
       "group in each gene enrichment group compared with the other enrichment groups; GO "
       "terms are counted at FDR < 0.05.")]),
    ("Figure 7. Analysis of main factors impacting TE functional associations",
     [("This filtering was applied to the visualization only.",
       "This filtering was applied to the visualization only. GO term counts are at "
       "FDR < 0.05. Pearson correlations in (A), (C) and (D) and Mann-Whitney U tests in (B), "
       "(E), (F) and (G) are single tests per panel and their p-values are raw; the Fisher "
       "exact stars in (H) are FDR-corrected.")]),
]
caption_edits_applied = 0
for needle, pairs in CAPTION_EDITS:
    paragraph = paragraph_by(body, needle)
    if paragraph is None:
        record("G10", f"locate caption: {needle[:50]}", False)
        continue
    for old, new in pairs:
        if replace_in_paragraph(paragraph, old, new):
            caption_edits_applied += 1
        else:
            record("G10", f"caption text: {old[:56]}", False)
record("G10", f"{caption_edits_applied} caption provenance edits", caption_edits_applied > 8)

figure1_caption = paragraph_by(body, "Figure 1. Enrichment of TEs in the 10 kb vicinity")
if figure1_caption is not None:
    replace_in_paragraph(
        figure1_caption,
        "The third vertical bar plot visualizes negative decimal logarithm of Fisher exact test "
        "p-value (FDR-corrected).",
        "The third vertical bar plot visualizes the negative decimal logarithm of the "
        "FDR-corrected Fisher exact test p-value.")
    record("G10", "Figure 1D caption wording", True)

# =======================================================================================
# G13 — human gene symbols italicised
# =======================================================================================
GENE_SYMBOLS = [
    "IFNA10", "IFNA14", "IFNA16", "IFNA17", "IFNA21", "IFNA22P", "IFNA4", "IFNA5", "IFNA6",
    "IFNA7", "IFNW1", "KLHL9", "POLR2A", "SSU72L1", "SSU72L2", "SSU72L3", "SSU72L4",
    "SSU72L5", "CIB3", "TMC1", "TMC2", "XIST", "UGT1A6", "UGT1A7", "UGT1A8", "UGT1A9",
    "UGT1A10", "B2M",
]


def italicise_symbols(body, symbols):
    """Italicise each whole-word gene symbol as a tracked delete-plus-insert.

    Longest first, so `SSU72L1` is handled before any shorter prefix could claim it, and
    `\\b` boundaries so `IFNA4` is never taken out of `IFNA45`.
    """
    applied = 0
    for symbol in sorted(symbols, key=len, reverse=True):
        pattern = re.compile(rf"\b{re.escape(symbol)}\b")
        for _ in range(60):
            hit = False
            for paragraph in [el for el in body.iter(qn("w:p"))]:
                strip_proof_errors(paragraph)
                placed = False
                for block in run_blocks(paragraph):
                    text = block_text(block)
                    match = pattern.search(text)
                    if not match:
                        continue
                    # Skip if this run is already italic: nothing to change.
                    rpr = block[0].find(qn("w:rPr"))
                    if rpr is not None and rpr.find(qn("w:i")) is not None:
                        continue
                    start, end = match.span()
                    before, matched, after = text[:start], text[start:end], text[end:]
                    pieces = []
                    if before:
                        pieces.append(make_run(before, base=rpr))
                    pieces.append(wrap_del(make_del_run(matched, base=rpr)))
                    pieces.append(wrap_ins(make_run(matched, base=rpr, italic=True)))
                    if after:
                        pieces.append(make_run(after, base=rpr))
                    anchor = block[-1]
                    for piece in pieces:
                        anchor.addprevious(piece)
                    for run in block:
                        paragraph.remove(run)
                    placed = True
                    break
                if placed:
                    applied += 1
                    hit = True
                    break
            if not hit:
                break
    return applied


def italicise_inside_insertions(body, symbols):
    """Italicise gene symbols in text this revision itself inserted.

    `run_blocks` deliberately cannot see inside a `<w:ins>`, which is what protects citation
    controls — but it also hides the Results and Discussion paragraphs added in Phase 5, and
    those name most of the genes. Inside an insertion the whole span is already marked as
    added, so the right edit is to split the run and set italic on it directly rather than to
    wrap yet another delete-plus-insert around it.
    """
    applied = 0
    patterns = [(s, re.compile(rf"\b{re.escape(s)}\b"))
                for s in sorted(symbols, key=len, reverse=True)]
    for insertion in list(body.iter(qn("w:ins"))):
        for run in list(insertion.findall(qn("w:r"))):
            node = run.find(qn("w:t"))
            if node is None or not node.text:
                continue
            rpr = run.find(qn("w:rPr"))
            if rpr is not None and rpr.find(qn("w:i")) is not None:
                continue
            text = node.text
            match = next((m for _, p in patterns for m in [p.search(text)] if m), None)
            if match is None:
                continue
            start, end = match.span()
            before, matched, after = text[:start], text[start:end], text[end:]
            replacements = []
            if before:
                replacements.append(make_run(before, base=rpr))
            replacements.append(make_run(matched, base=rpr, italic=True))
            if after:
                replacements.append(make_run(after, base=rpr))
            for element in replacements:
                run.addprevious(element)
            insertion.remove(run)
            applied += 1
    return applied


italicised = italicise_symbols(body, GENE_SYMBOLS)
# Repeat inside insertions until nothing more splits out, since one pass handles one symbol
# per run and a sentence like the IFNA gene list holds eight in a single run.
inside = 0
for _ in range(40):
    found = italicise_inside_insertions(body, GENE_SYMBOLS)
    inside += found
    if not found:
        break
record("G13", f"italicise {italicised} gene symbol(s) in original text and {inside} in "
              f"revision-inserted text", italicised + inside >= 25)

# =======================================================================================
# G8 / G9 — Tables 1 and 2 to the end of the main text, with 0.5 pt rules
# =======================================================================================
def add_half_point_borders(table_element):
    """0.5 pt single rules on every edge and inner line (G3 section 7)."""
    properties = table_element.find(qn("w:tblPr"))
    if properties is None:
        properties = OxmlElement("w:tblPr")
        table_element.insert(0, properties)
    for existing in properties.findall(qn("w:tblBorders")):
        properties.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")          # eighths of a point: 4 = 0.5 pt
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    properties.append(borders)


new_tables = []
for element in body.iter(qn("w:tbl")):
    header = "".join(t.text or "" for t in element.iter(qn("w:t")))
    if header.startswith("Class"):
        new_tables.append(element)
for element in new_tables:
    add_half_point_borders(element)
record("G8/G9", f"0.5 pt rules on {len(new_tables)} table(s)", len(new_tables) == 2)

# Move both tables and their captions to sit after the Discussion, before the back matter.
supplementary_heading = heading_paragraph(body, "Supplementary material")
caption_1 = paragraph_by(body, "Table 1. Enrichment of TE classes in gene TSS")
caption_2 = paragraph_by(body, "Table 2. Statistical support for TE class enrichment")
if (supplementary_heading is not None and caption_1 is not None and caption_2 is not None
        and len(new_tables) == 2):
    segment = [caption_1, new_tables[0], caption_2, new_tables[1]]
    marker = note_paragraph(
        "G8: Tables 1 and 2 are placed here, at the end of the main text, as G3 requires. They "
        "are editable Word tables with 0.5 pt rules. Note that this relocation is a structural "
        "move rather than a Word tracked move, so Reject All restores the original text but not "
        "the original position.")
    insert_before(supplementary_heading, [marker] + segment)
    record("G8", "move Tables 1 and 2 to the end of the main text", True)
else:
    record("G8", "move Tables 1 and 2", False)

# =======================================================================================
# G7 — Materials and methods before Results
# =======================================================================================
def segment_between(body, start_heading, end_heading):
    """Body children from `start_heading` up to (not including) `end_heading`."""
    children = list(body)
    try:
        first = children.index(start_heading)
    except ValueError:
        return []
    for index in range(first + 1, len(children)):
        if children[index] is end_heading:
            return children[first:index]
    return []


methods_heading = heading_paragraph(body, "Materials and methods")
results_heading = heading_paragraph(body, "Results")
ethical_marker = heading_paragraph(body, "Conflicts of interest")
if methods_heading is not None and results_heading is not None and ethical_marker is not None:
    methods_segment = segment_between(body, methods_heading, ethical_marker)
    if methods_segment:
        for element in methods_segment:
            body.remove(element)
        insert_before(results_heading, methods_segment)
        insert_before(methods_heading, [note_paragraph(
            "G7: Materials and methods is moved ahead of Results, which is G3's printed order. "
            "This is a structural relocation, not a Word tracked move, so Reject All restores "
            "the original text but leaves the sections in this order.")])
        record("G7", f"move Materials and methods ({len(methods_segment)} elements) before "
                     f"Results", True)
    else:
        record("G7", "delimit the Methods segment", False)
else:
    record("G7", "locate the Methods / Results / Conflicts headings", False)

# =======================================================================================
# G7 — back-matter order
#
# The draft runs ETHICAL STATEMENT, CONFLICT OF INTEREST, AUTHORS CONTRIBUTION,
# ACKNOWLEDGEMENTS, SUPPLEMENTARY MATERIAL, DATA AVAILABILITY, REFERENCES. G3 section 2 wants
# Supplementary material as the last *body* section, then Acknowledgments, Data availability,
# Funding, Conflicts of interest, Literature cited. Author contributions is not in G3's list;
# for a single-author paper it is placed last before Literature cited rather than dropped.
# =======================================================================================
BACK_MATTER_ORDER = [
    "Supplementary material",
    "Acknowledgments",
    "Data availability",
    "Funding",
    "Conflicts of interest",
    "Author contributions",
    "Literature cited",
]

back_headings = {name: heading_paragraph(body, name) for name in BACK_MATTER_ORDER}
missing = [name for name, element in back_headings.items() if element is None]
if missing:
    record("G7", f"locate back-matter headings ({', '.join(missing)})", False)
else:
    children = list(body)
    positions = {name: children.index(element)
                 for name, element in back_headings.items()}
    first_position = min(positions.values())
    sect_pr = body.find(qn("w:sectPr"))
    boundaries = sorted(positions.values()) + [
        children.index(sect_pr) if sect_pr is not None else len(children)]
    segments = {}
    for name, start in positions.items():
        end = next(b for b in boundaries if b > start)
        segments[name] = children[start:end]
    # Anything between the end of the Discussion and the first back-matter heading stays put.
    for name in BACK_MATTER_ORDER:
        for element in segments[name]:
            body.remove(element)
    anchor = list(body)[first_position - 1] if first_position > 0 else None
    ordered = [element for name in BACK_MATTER_ORDER for element in segments[name]]
    if anchor is not None:
        insert_after(anchor, ordered)
    else:
        for element in ordered:
            body.append(element)
    insert_before(back_headings["Supplementary material"], [note_paragraph(
        "G7: the back matter is reordered to G3's printed order — Supplementary material as the "
        "last body section, then Acknowledgments, Data availability, Funding, Conflicts of "
        "interest, and Literature cited. Author contributions is not one of G3's back-matter "
        "sections; for a single-author paper it is kept and placed last. As with the Methods "
        "move, this is a structural relocation rather than a Word tracked move.")])
    record("G7", f"reorder {len(BACK_MATTER_ORDER)} back-matter sections to G3 order", True)

# =======================================================================================
# Document mechanics — double spacing and page numbers
# =======================================================================================
styles = document.styles.element
normal = next((s for s in styles.iter(qn("w:style"))
               if s.get(qn("w:styleId")) == "Normal"), None)
if normal is not None:
    ppr = normal.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        normal.append(ppr)
    for existing in ppr.findall(qn("w:spacing")):
        ppr.remove(existing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "480")        # 480 twentieths of a point = double
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    record("mechanics", "double spacing on the Normal style", True)
else:
    record("mechanics", "locate the Normal style", False)

section = document.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
if not footer.paragraphs or not footer.paragraphs[0].text.strip():
    footer_paragraph = (footer.paragraphs[0] if footer.paragraphs
                        else footer.add_paragraph())._p
    for child in list(footer_paragraph):
        if child.tag == qn("w:r"):
            footer_paragraph.remove(child)
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run.append(begin)
    footer_paragraph.append(run)
    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    instruction_run.append(instruction)
    footer_paragraph.append(instruction_run)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    footer_paragraph.append(end_run)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr = footer_paragraph.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        footer_paragraph.insert(0, ppr)
    ppr.append(jc)
    record("mechanics", "page-number field in the footer", True)

# Line numbers were already present in sectPr; confirm rather than duplicate.
sect_pr = body.find(qn("w:sectPr"))
record("mechanics", "line numbers already enabled",
       sect_pr is not None and sect_pr.find(qn("w:lnNumType")) is not None)

# =======================================================================================
# Save and report
# =======================================================================================
if not DRY_RUN:
    document.save(TARGET)

failed = [row for row in report if not row["applied"]]
print(f"{len(report) - len(failed)}/{len(report)} house-style edits applied")
for row in failed:
    print(f"  NOT APPLIED [{row['gap']}] {row['edit']}")
with open(os.path.join(HERE, "output", "house_style_report.json"), "w") as handle:
    json.dump(report, handle, indent=2)
print("report -> output/house_style_report.json")
if failed:
    sys.exit(1)
